from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "提交材料"


def main():
    print(f"doc_dir_exists={DOC_DIR.exists()}")
    for path in sorted(DOC_DIR.glob("*.docx")):
        doc = Document(path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        refs = sum(1 for p in paragraphs if p.startswith("["))
        print(
            f"{path.name}\tparagraphs={len(paragraphs)}\t"
            f"tables={len(doc.tables)}\treferences={refs}\t"
            f"title={paragraphs[0] if paragraphs else ''}"
        )


if __name__ == "__main__":
    main()
