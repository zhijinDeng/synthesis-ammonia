from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "提交材料"
OUT.mkdir(exist_ok=True)

FONT_CN = "Microsoft YaHei"
FONT_EN = "Calibri"
BLUE = RGBColor(31, 78, 121)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


REFERENCES = [
    "IEA. Ammonia Technology Roadmap. https://www.iea.org/reports/ammonia-technology-roadmap",
    "IEA. Ammonia Technology Roadmap: Executive Summary. https://www.iea.org/reports/ammonia-technology-roadmap/executive-summary",
    "成都云图控股股份有限公司. 2025 年年度报告摘要. 巨潮资讯. https://static.cninfo.com.cn/finalpage/2026-04-15/1225100697.PDF",
    "成都云图控股股份有限公司. 投资者关系活动记录表. 巨潮资讯. https://static.cninfo.com.cn/finalpage/2026-04-16/1225109829.PDF",
    "云图控股研究报告. 合成氨、磷矿等新产能落地或助力公司成长. https://www.wintrueholding.com/static/upload/file/20260417/1776394254161882.pdf",
    "Kong et al. Nonlinear Model Predictive Control of Flexible Ammonia Production. 2024 preprint. https://qizh.cems.umn.edu/sites/qizh.cems.umn.edu/files/2024-04/Kong_et_al-2024-preprint-NMPC_ammonia_production.pdf",
    "ACS Industrial & Engineering Chemistry Research. Dynamic Simulation and Optimization for Load Regulation of the Haber-Bosch Process. https://pubs.acs.org/doi/10.1021/acs.iecr.4c02410",
    "AIChE 2025 Annual Meeting. Scheduling and Control of Green Ammonia Plants. https://aiche.confex.com/aiche/2025/meetingapp.cgi/Paper/710466",
    "中国信息通信研究院、华为等. 工业与AI融合应用指南. https://www-file.huawei.com/dam/asset/view/dda7ee674ced41d4bb989b365ea0a028.pdf",
    "HG/T 6346—2025. 石化和化工行业数字化转型成熟度模型与评估相关资料. https://gxt.gansu.gov.cn/gxt/c126826/202511/174237249/files/99bb56ef71d347248ae58125b2e12aee.pdf",
    "成都云图控股股份有限公司. 关于全资子公司 70 万吨合成氨项目建成试生产的公告. 巨潮资讯. https://static.cninfo.com.cn/finalpage/2026-03-21/1225020927.PDF",
    "AQ/T 3017—2008. 合成氨生产企业安全标准化实施指南. 应急管理部. https://www.mem.gov.cn/fw/flfgbz/bz/bzwb/201301/P020190327398204066876.pdf",
    "合成氨行业规范条件. 工业和信息化部相关资料. https://file.smejs.cn/group1/M00/05/1F/rBIAAWUvo0OAbXwqAAUaBh6SZ0U730.pdf",
    "国家发展改革委. 合成氨行业节能降碳改造升级实施指南解读. https://www.ndrc.gov.cn/xwdt/ztzl/ghnhyjnjdgzsj/zjgd/202208/t20220829_1334056.html",
    "应急管理部办公厅. “工业互联网+危化安全生产”试点建设方案. https://www.mem.gov.cn/gk/zfxxgkpt/fdzdgknr/202104/t20210407_382882.shtml",
    "GB 17681—2024. 危险化学品重大危险源安全监控技术规范相关资料. https://sjj.ganzhou.gov.cn/c101840/202506/e42a2cfd5c894799b899a30fcb927f06/files/b133fc57bfa64bd3a7beeee6ce6ca38e.pdf",
]

PART1_SUBMISSION = (
    "这道题的核心不是再造一个看板，而是在合成氨硬件成本已接近极限后，用AI挖掘运行管理的软件降本空间。"
    "云图应城合成氨装置连接尿浆、复合肥、联碱和液氨外售，调度同时受氢氮比、合成塔温升、循环压缩机、液氨库存、下游订单和能耗约束。"
    "企业想看到的能力，是把行业资料、现场机理、专家经验和实时数据合成一个会寻优、会指挥、会复盘、会学习的数字调度员，而不是泛泛讲大模型。"
)

PART2_SUBMISSION = (
    "我设计“合成氨AI调控师工作台”，先以影子运行接入MES日计划、ERP订单交期、DCS关键点位、液氨罐区、设备点检和能源价格，形成班次事实表。"
    "系统用预测模型滚动判断下游用氨、库存水位、吨氨能耗和设备风险；用机理约束校验氢氮比、合成塔床层温升、循环压缩机负荷、罐区压力、安全库存和负荷升降速率；"
    "再由优化器生成稳氨、保供、护机三案，给出推荐负荷、下游分配、成本收益、风险边界和需人工确认事项。"
    "调度员采纳后，平台跟踪实际负荷、氢氮比偏差、液氨库存变化、订单完成和吨氨成本，自动生成班后复盘。"
    "每次采纳或未采纳原因都会进入合成负荷指令库、异常经验库和专家规则库，用于周度校准。"
    "落地按30/60/90天推进：先打通数据口径，再影子对比人工调度，最后小闭环验收。"
    "价值以高优订单满足率、单位氨能耗、库存占用、调度耗时和异常预警提前量量化；低置信、数据延迟或安环红线接近时自动降级人工调度。"
)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(row.cells[idx])


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = FONT_CN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def setup_doc(title, subtitle):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, RGBColor(31, 77, 120))]:
        st = styles[name]
        st.font.name = FONT_CN
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.text = "合成氨 AI 生产调度专家 | 提交材料"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], size=9, color=RGBColor(90, 90, 90))

    footer = section.footer.paragraphs[0]
    footer.text = "仅用于参赛方案与试点论证"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.runs[0], size=9, color=RGBColor(120, 120, 120))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=22, bold=True, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(subtitle)
    set_run_font(r, size=11, color=RGBColor(80, 80, 80))

    return doc


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        r = cell.paragraphs[0].add_run(header)
        set_run_font(r, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            r = cells[idx].paragraphs[0].add_run(str(value))
            set_run_font(r, size=9.5)
    set_table_width(table, widths)
    return table


def add_references(doc):
    doc.add_heading("参考文献与资料来源", level=1)
    for idx, ref in enumerate(REFERENCES, start=1):
        add_para(doc, f"[{idx}] {ref}")


def save(doc, name):
    path = OUT / name
    doc.save(path)
    return path


def build_submission_form_doc():
    doc = setup_doc("报名表填写文本", "按截图要求整理：文本框、附件、链接与队伍信息")
    doc.add_heading("1. 开题报告 Part 1：命题前置分析与洞察", level=1)
    add_para(doc, f"建议粘贴文本（{len(PART1_SUBMISSION)} 字）：", bold_prefix="建议粘贴文本")
    add_para(doc, PART1_SUBMISSION)

    doc.add_heading("2. 开题报告 Part 2：整体解决方案设计", level=1)
    add_para(doc, f"建议粘贴文本（{len(PART2_SUBMISSION)} 字）：", bold_prefix="建议粘贴文本")
    add_para(doc, PART2_SUBMISSION)

    doc.add_heading("3. 团队补充材料（附件格式）", level=1)
    add_table(
        doc,
        ["建议上传", "作用", "对应文件"],
        [
            ["开题报告", "直接回应两个必填文本框，并保留出题意图、解题思路、文献依据", "01_开题报告.docx"],
            ["整体解决方案书", "证明方案不是概念页，而是有架构、接口、流程、验收和停用边界", "02_整体解决方案书.docx"],
            ["调控师操作手册", "从企业班组视角说明如何用、谁确认、何时降级、如何复盘", "03_调控师操作手册.docx"],
            ["参考文献与数据依据", "展示行业报告、公司资料、控制优化论文和化工数字化标准依据", "04_参考文献与数据依据.docx"],
            ["企业评审优化清单", "站在评审视角回答安全、收益、责任、数据、推广和创新辨识度", "05_企业评审视角优化清单.docx"],
        ],
        [1.5, 2.7, 2.2],
    )

    doc.add_heading("4. 团队补充材料（链接格式）", level=1)
    add_para(doc, "建议填写 GitHub 项目链接：https://github.com/zhijinDeng/synthesis-ammonia")
    add_para(doc, "链接说明：仓库包含可交互演示页面、合成氨装置数据模型、研究简报、架构说明、优化路线、Word 提交材料生成脚本和提交附件。")

    doc.add_heading("5. 队伍参赛人数", level=1)
    add_para(doc, "建议选择：1 人。")
    add_para(doc, "说明：若平台要求 1-3 人，可按实际报名人数填写；当前材料以单人完整方案包组织，突出个人对行业、AI、工程落地和文档交付的综合能力。")
    return save(doc, "00_报名表填写文本.docx")


def build_opening_report():
    doc = setup_doc("开题报告", "命题：如何用 AI 打造合成氨生产调度专家")
    doc.add_heading("报名表可直接提交正文", level=1)
    doc.add_heading(f"Part 1：命题前置分析与洞察（{len(PART1_SUBMISSION)} 字）", level=2)
    add_para(doc, PART1_SUBMISSION)
    doc.add_heading(f"Part 2：整体解决方案设计（{len(PART2_SUBMISSION)} 字）", level=2)
    add_para(doc, PART2_SUBMISSION)
    doc.add_heading("补充展开：为什么这样答更贴合企业", level=1)
    add_para(
        doc,
        "截图要求强调“跳出基础信息、搜集行业报告、竞品案例、用户真实反馈并输出理解与思考”。因此开题报告不能写成技术综述，而要先把试题的产业语境讲清楚：云图的合成氨不只是单套装置稳产，而是连接尿浆、复合肥、联碱、液氨外售和氮肥自给率的经营枢纽。"
    )
    add_para(
        doc,
        "整体方案也不能只写算法名词。评委真正会关注五件事：架构是否完整、创新是否区别于普通看板、业务价值是否可量化、试点是否能落地、未来是否可推广。因此本方案采用“影子运行先行、人机确认闭环、DCS/SIS 守底线、采纳效果进知识库”的路径，把 AI 定位为合成氨调度员的专业副驾驶。"
    )
    doc.add_heading("三大挑战与回答", level=1)
    add_table(
        doc,
        ["试题挑战", "方案回答", "企业可见成果"],
        [
            ["调度寻优难", "吨氨成本雷达 + 稳氨/保供/护机三案 + 合成回路约束孪生", "看到每次降本来自原料/蒸汽、液氨库存、压缩机、合成塔还是负荷偏差"],
            ["指挥效率低", "合成氨调控师指令窗口 + 交接摘要 + 执行监督", "缩短班前会准备和合成负荷调整确认时间"],
            ["知识传承难", "合成负荷指令库 + 氢氮比/床层温升异常库 + 未采纳标签", "把专家经验变成可检索、可复盘、可迭代的合成氨知识资产"],
        ],
        [1.2, 2.6, 2.7],
    )
    doc.add_heading("出题意图判断", level=1)
    add_bullets(
        doc,
        [
            "业务洞察：能否理解云图控股合成氨项目与复合肥、尿浆、联碱产业链的协同关系。",
            "工业现场理解：能否意识到合成氨高温高压、连续流程、安全边界和设备健康的重要性。",
            "AI 工程理解：能否把预测、优化、数字孪生、知识库和人机确认组合起来，而不是泛泛讲大模型。",
            "落地能力：能否提出接口、岗位、验收、降级和收益核算，而不是只有概念架构。",
        ],
    )
    add_references(doc)
    return save(doc, "01_开题报告.docx")


def build_solution_doc():
    doc = setup_doc("整体解决方案书", "合成氨 AI 调控师工作台：从影子运行到人机协同闭环")
    doc.add_heading("1. 方案定位", level=1)
    add_para(doc, "本方案把合成氨装置视为多约束、多目标、滚动优化的生产经营系统。AI 的角色不是替代班长，而是担任“调控师副驾驶”：在安全边界内汇总数据、识别约束、生成方案、解释风险，并把采纳与偏差沉淀为知识。")
    doc.add_heading("2. 首批试点范围", level=1)
    add_table(
        doc,
        ["范围", "先做", "暂不做"],
        [
            ["装置", "合成氨、液氨库存、尿浆/复合肥需求联动", "全厂所有装置一次性纳入"],
            ["数据", "MES、ERP、DCS 摘要点、罐区、设备健康评分", "全量秒级点位接入"],
            ["执行", "排产建议、交接摘要、复盘记录", "自动开停车或直接写控制参数"],
            ["收益", "采纳方案的实际差额收益", "把市场行情自然波动算作 AI 收益"],
        ],
        [1.2, 2.7, 2.6],
    )
    doc.add_heading("3. 合成氨超级数字员工能力矩阵", level=1)
    add_table(
        doc,
        ["能力", "功能模块", "运行闭环"],
        [
            ["实时感知", "氢氮比、合成塔温升、循环压缩机、液氨库存、下游用氨统一采集", "形成合成氨班次事实表"],
            ["吨氨寻优", "吨氨成本雷达、稳氨/保供/护机三案、合成回路约束", "识别隐藏成本白银和真金"],
            ["敏捷指挥", "合成氨调控师指令窗口、交接摘要、审批链", "人工确认后下达合成负荷计划"],
            ["执行监督", "合成负荷偏差、氢氮比偏差、液氨库存、吨氨能耗结果跟踪", "自动生成复盘样本"],
            ["知识传承", "合成负荷指令库、异常经验库、未采纳原因标签", "沉淀合成氨专家知识和场景规则"],
            ["自我迭代", "周度校准、规则修订、高风险样本复盘", "根据实际生产持续升级"],
        ],
        [1.1, 2.7, 2.7],
    )
    doc.add_heading("4. 系统架构", level=1)
    add_bullets(
        doc,
        [
            "数据层：形成班次事实表，包含计划、订单、库存、负荷、能耗、设备、异常和审批记录。",
            "预测层：滚动预测订单压力、库存水位、能源窗口、设备健康和市场波动。",
            "约束层：把合成塔、压缩机、罐区、蒸汽平衡、安全库存、环保指标作为硬边界。",
            "优化层：输出稳态、冲刺、保守三套方案，并进行订单、库存、能耗和风险权衡。",
            "解释层：生成班前会摘要、风险解释、备选方案和班后复盘要点。",
            "治理层：记录输入快照、模型版本、审批人、采纳原因、执行偏差和回滚条件。",
        ],
    )
    doc.add_heading("5. 接口与数据清单", level=1)
    add_table(
        doc,
        ["系统", "首批字段", "频率", "用途"],
        [
            ["MES", "日计划、班次产量、执行偏差、偏差原因", "班次/小时", "计划与实际对比"],
            ["ERP", "订单、交期、价格口径、客户优先级", "日/订单变更", "订单优先级与延期成本"],
            ["DCS historian", "负荷、温度、压力、流量、电耗、蒸汽摘要点", "5-15 分钟聚合", "装置边界判断"],
            ["罐区系统", "液氨库存、罐区压力、装车窗口", "分钟/小时", "库存和外售约束"],
            ["EAM/点检", "设备健康评分、检修计划、异常工单", "日/事件", "设备风险约束"],
        ],
        [1.2, 2.2, 1.2, 1.9],
    )
    doc.add_heading("6. 30/60/90 天落地路线", level=1)
    add_table(
        doc,
        ["时间", "目标", "交付物", "验收口径"],
        [
            ["30 天", "跑通班次事实表", "接口清单、数据质量报告、口径确认表", "关键字段完整率达标，调度确认口径"],
            ["60 天", "影子运行", "每班三套建议、未采纳原因记录", "覆盖主要场景，调度员愿意看"],
            ["90 天", "小闭环复盘", "收益归因、规则库、上线边界", "至少一项指标稳定改善或明确终止原因"],
        ],
        [0.9, 1.5, 2.2, 1.9],
    )
    doc.add_heading("7. 验收和停用", level=1)
    add_bullets(
        doc,
        [
            "通过条件：高优订单满足率不低于人工排产，且单位氨能耗、库存占用、调度耗时或异常预警至少一项改善。",
            "停用条件：关键数据延迟超过 15 分钟、核心点位缺失、模型连续两天偏差超阈值或安环红线接近且无法解释。",
            "推广条件：完成一个月影子运行、三次异常场景复盘和一份采纳/不采纳原因清单。",
        ],
    )
    add_references(doc)
    return save(doc, "02_整体解决方案书.docx")


def build_operator_manual():
    doc = setup_doc("调控师操作手册", "合成氨 AI 调控师工作台：班组使用与异常处置")
    doc.add_heading("1. 调控师角色", level=1)
    add_para(doc, "调控师不是替代班长的自动控制系统，而是帮助当班人员在订单、库存、能耗、设备和安全边界之间快速形成可解释方案的副驾驶。所有高风险动作必须保留人工确认。")
    doc.add_heading("2. 当班流程", level=1)
    add_table(
        doc,
        ["环节", "操作", "输出", "责任人"],
        [
            ["班前 30 分钟", "刷新订单、库存、设备健康、能源窗口", "稳态/冲刺/保守三套方案", "调度员"],
            ["班前会 10 分钟", "查看差异、红线、需确认动作", "交接摘要与确认项", "班长"],
            ["班中偏差", "库存、设备、订单超阈值时重算", "重算记录，不自动改控制参数", "班长/调度"],
            ["班后 5 分钟", "记录采纳与否、实际负荷、能耗、异常原因", "复盘样本", "班长"],
        ],
        [1.2, 2.1, 2.2, 1.0],
    )
    doc.add_heading("3. 典型调控场景", level=1)
    add_table(
        doc,
        ["场景", "触发条件", "推荐动作", "不得越界"],
        [
            ["订单冲刺", "下游订单压力高，库存偏低", "提高合成负荷，尿浆/复合肥优先，外售降级", "不得低于安全库存"],
            ["设备保守", "压缩机或合成塔健康评分下降", "降低负荷上限，插入点检窗口", "不得为追产突破设备边界"],
            ["能源错峰", "高价能源窗口或蒸汽紧张", "高价窗口压产，低价窗口补库存", "不得影响刚性交付和安全库存"],
            ["低置信", "数据缺失、模型偏差或市场异常", "退回规则建议和人工流程", "不得自动回写 MES"],
        ],
        [1.1, 1.8, 2.4, 1.2],
    )
    doc.add_heading("4. 交接班摘要模板", level=1)
    add_bullets(
        doc,
        [
            "推荐负荷：__%；推荐模式：稳态/冲刺/保守。",
            "优先订单：尿浆/复合肥/液氨外售/联碱配套。",
            "关键约束：库存、压缩机、合成塔、罐区、能源窗口、环保指标。",
            "审批要求：班长确认/班长 + 调度主管双确认。",
            "未采纳原因：安全边界、设备风险、订单变化、数据不可信、经验判断。",
        ],
    )
    doc.add_heading("5. 调控师判断原则", level=1)
    add_numbered(
        doc,
        [
            "先安全后收益：任何收益测算都不能覆盖安环和设备硬边界。",
            "先小闭环后推广：一个班次链路跑通，再谈全厂协同。",
            "先解释后执行：无法解释触发约束的建议不得采纳。",
            "先复盘后迭代：未采纳原因比采纳原因更能提升模型。",
        ],
    )
    return save(doc, "03_调控师操作手册.docx")


def build_reference_doc():
    doc = setup_doc("参考文献与数据依据", "合成氨 AI 调度方案证据链")
    doc.add_heading("1. 关键事实", level=1)
    add_bullets(
        doc,
        [
            "氨是氮肥生产的重要基础原料，氨行业具有显著能源消耗和碳排放特征[1][2]。",
            "云图控股应城基地合成氨项目与尿浆、复合肥等产品具有产业链协同价值[3][4][5]。",
            "柔性合成氨和 Haber-Bosch 负荷调节研究说明，连续装置的动态负荷优化具有研究基础，但必须受机理约束和现场规则限制[6][7][8]。",
            "工业 AI 先进架构可作为支撑方法，但必须落到合成氨的氢氮比、合成塔、循环压缩机、液氨库存等对象上[9]。",
            "石化和化工数字化成熟度要求强调生产平衡经验库、调度指令库、异常处置经验库、自动生成调度指令并跟踪执行结果，可支撑合成氨知识库闭环设计[10]。",
        ],
    )
    doc.add_heading("2. 资料如何支撑方案", level=1)
    add_table(
        doc,
        ["资料类型", "用于支撑", "在方案中的体现"],
        [
            ["行业路线图", "能耗、碳排、氨产业位置", "能耗优化、碳强度下降、保守收益口径"],
            ["云图材料", "合成氨项目与下游协同", "尿浆、复合肥、联碱、液氨外售统一调度"],
            ["控制与优化论文", "柔性负荷和 MPC 可行性", "MPC + 规则库 + 人机确认"],
            ["工业调度研究", "数字孪生和调度控制趋势", "影子运行、模型治理、异常降级"],
            ["工业 AI 指南", "智能体、数字孪生、知识闭环", "作为合成氨 AI 调控师的技术支撑，而非泛化主题"],
            ["化工成熟度要求", "生产平衡经验库、调度指令库、执行反馈", "支撑合成氨执行监督闭环和知识库沉淀"],
        ],
        [1.4, 2.2, 2.9],
    )
    add_references(doc)
    return save(doc, "04_参考文献与数据依据.docx")


def build_review_upgrade_doc():
    doc = setup_doc("企业评审视角优化清单", "把合成氨 AI 调控师从演示方案升级为企业试点提案")
    doc.add_heading("1. 企业评审真正会看什么", level=1)
    add_table(
        doc,
        ["评审维度", "常见扣分点", "本方案强化点"],
        [
            ["业务贴合", "把合成氨写成普通生产看板，忽视下游消纳", "围绕氨合成负荷、液氨库存、尿浆、复合肥、联碱和液氨外售做协同优化"],
            ["安全可信", "只谈降本，不谈高温高压、重大危险源和审批边界", "把合成塔温升、循环压缩机、罐区压力、安全库存、安环红线作为硬约束"],
            ["现场可用", "页面漂亮但班长不知道怎么用", "新增调控工作说明、专家审核关注点、班前会摘要、飞书卡片和审批草稿"],
            ["收益可信", "泛泛写降本增效，没有核算边界", "只统计被采纳方案的实际差额收益，剔除行情自然波动和外部检修影响"],
            ["知识传承", "把知识库写成口号", "记录采纳/未采纳原因、输入快照、审批人、实际偏差和班长经验判断"],
            ["推广复制", "一上来做全厂大平台", "先 30/60/90 天影子运行，再小闭环，再复制到多装置协同"],
        ],
        [1.1, 2.4, 3.0],
    )

    doc.add_heading("2. 新增突出创新点", level=1)
    add_bullets(
        doc,
        [
            "专家审核友好：页面不再由参赛方自我打分，而是用“调控工作说明”和“专家审核关注点”呈现待确认事项、审批责任和试点条件。",
            "反事实收益归因：同一班次保留原计划、AI 建议、采纳结果和实际偏差，用反事实对比减少“把行情算成 AI 收益”的争议。",
            "三重护栏：工艺硬约束、安环合规护栏、飞书审批留痕同时生效，证明 AI 不是绕开控制系统。",
            "未采纳原因资产化：把安全边界、设备风险、订单变化、数据不可信、经验判断等未采纳原因结构化，反向训练数字员工。",
            "产业链协同视角：氨不是孤立产品，而是尿浆、复合肥、联碱、液氨外售和氮肥自给率的枢纽原料。",
            "轻集成落地：飞书负责通知、审批、复盘和协同，MES/DCS/ERP 仍保留原有系统边界，降低企业试点阻力。",
        ],
    )

    doc.add_heading("3. 评委可能追问与回答", level=1)
    add_table(
        doc,
        ["追问", "回答口径"],
        [
            ["AI 会不会影响安全生产？", "不会。第一阶段只做影子建议、飞书审批和复盘，不直接写 DCS/SIS，不自动开停车；安环红线接近时自动降级。"],
            ["数据不准怎么办？", "先做数据就绪度评分。关键点位缺失、延迟超过 15 分钟或口径不一致时，优化建议停用，只保留人工规则提示。"],
            ["为什么需要飞书？", "飞书不是控制系统，而是班组协同入口。它让建议、审批、采纳、驳回、复盘都有责任人和时间戳。"],
            ["收益怎么证明？", "用历史回放和影子运行做对照，只认被采纳且实际执行的方案收益，并按能源、库存、订单、停车风险分类归因。"],
            ["为什么方案不像外行？", "方案从合成回路、循环压缩机、合成塔床层温升、氢氮比、液氨罐区和下游用氨出发，而不是泛泛写 AI 大模型。"],
            ["能否推广？", "先沉淀合成氨约束模板、审批模板和知识库，再复制到复合肥、联碱、黄磷、磷酸铁等多装置协同场景。"],
        ],
        [1.9, 4.6],
    )

    doc.add_heading("4. 推荐参赛表达", level=1)
    add_para(
        doc,
        "如果评委只看平台，应该先看到一个会说“我先守安全边界，再找吨氨成本”的合成氨调控师；如果评委打开附件，应该看到方案不是一个 AI 页面，而是一套能进入企业试点的运行管理系统。"
    )
    add_para(
        doc,
        "最终表达建议：不把 AI 描述成全能接管者，而描述成“懂合成氨、懂班组、懂审批、懂复盘的数字调控师”。它的价值在于把专家经验从人脑、微信群和交接班口头记录中沉淀成可复盘、可审计、可迭代的企业知识资产。"
    )

    doc.add_heading("5. 额外证据补强", level=1)
    add_bullets(
        doc,
        [
            "云图控股项目公告和年报用于支撑“70 万吨合成氨项目与下游复合肥、尿浆协同”的业务背景。",
            "合成氨安全标准化实施指南用于支撑企业安全生产和现场管理边界。",
            "合成氨行业规范条件和节能降碳资料用于支撑能耗、碳排、装置效率和改造升级价值。",
            "工业互联网+危化安全生产和重大危险源监控资料用于支撑飞书审批、事件留痕和安全监控闭环。",
        ],
    )
    add_references(doc)
    return save(doc, "05_企业评审视角优化清单.docx")


if __name__ == "__main__":
    paths = [
        build_submission_form_doc(),
        build_opening_report(),
        build_solution_doc(),
        build_operator_manual(),
        build_reference_doc(),
        build_review_upgrade_doc(),
    ]
    for path in paths:
        print(path)
