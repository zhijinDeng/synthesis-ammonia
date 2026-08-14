from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MATERIALS = ROOT / "提交材料"

FONT_CN = "Microsoft YaHei"
FONT_EN = "Calibri"
BLUE = RGBColor(24, 78, 119)
DARK = RGBColor(13, 31, 48)
GRAY = RGBColor(90, 105, 118)
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F4F6F8"
LIGHT_GREEN = "EAF7EF"


REFERENCES = [
    "IEA. Ammonia Technology Roadmap. International Energy Agency. https://www.iea.org/reports/ammonia-technology-roadmap",
    "成都云图控股股份有限公司. 2025年年度报告摘要. 巨潮资讯. https://static.cninfo.com.cn/finalpage/2026-04-15/1225100697.PDF",
    "成都云图控股股份有限公司. 关于全资子公司70万吨合成氨项目建成试生产的公告. 巨潮资讯. https://static.cninfo.com.cn/finalpage/2026-03-21/1225020927.PDF",
    "成都云图控股股份有限公司. 投资者关系活动记录表. 巨潮资讯. https://static.cninfo.com.cn/finalpage/2026-04-16/1225109829.PDF",
    "Kong et al. Nonlinear Model Predictive Control of Flexible Ammonia Production. 2024 preprint. https://qizh.cems.umn.edu/sites/qizh.cems.umn.edu/files/2024-04/Kong_et_al-2024-preprint-NMPC_ammonia_production.pdf",
    "Dynamic Simulation and Optimization for Load Regulation of the Haber-Bosch Process. Industrial & Engineering Chemistry Research. https://pubs.acs.org/doi/10.1021/acs.iecr.4c02410",
    "AQ/T 3017-2008. 合成氨生产企业安全标准化实施指南. 应急管理部. https://www.mem.gov.cn/fw/flfgbz/bz/bzwb/201301/P020190327398204066876.pdf",
    "国家发展改革委. 合成氨行业节能降碳改造升级实施指南解读. https://www.ndrc.gov.cn/xwdt/ztzl/ghnhyjnjdgzsj/zjgd/202208/t20220829_1334056.html",
    "应急管理部办公厅. 工业互联网+危化安全生产试点建设方案. https://www.mem.gov.cn/gk/zfxxgkpt/fdzdgknr/202104/t20210407_382882.shtml",
    "International Society of Automation. ISA-95 Series of Standards: Enterprise-Control System Integration. https://www.isa.org/standards-and-publications/isa-standards/isa-95-standard",
    "Aspen Technology. Aspen GDOT. https://www.aspentech.com/en/products/msc/aspen-gdot",
    "Honeywell. Honeywell Forge Performance+ for Industrials | Production Intelligence. https://process.honeywell.com/us/en/products/industrial-software/operational-excellence/honeywell-forge-production-intelligence",
    "International Society of Automation. ISA-18 Series of Standards. https://www.isa.org/standards-and-publications/isa-standards/isa-18-series-of-standards",
    "飞书开放平台. 任务概述（Task v2）. https://open.feishu.cn/document/task-v2/overview",
    "飞书开放平台. 开放平台概述（含飞书智能伙伴搭建平台Aily）. https://open.feishu.cn/document/client-docs/intro?lang=zh-CN",
    "飞书开放平台. 执行数据知识问答（Aily v1）. https://open.feishu.cn/document/uAjLw4CM/ukTMukTMukTM/aily-v1/data-knowledge",
]


PART1 = (
    "我们将课题界定为班次级生产调度问题：在安全和连续生产约束下，综合液氨库存、下游需求、订单、价格、"
    "公用工程和设备状态，确定液氨分配、装置负荷、外采安排及必要的停运评估。平台连接现有系统的数据和班组流程，"
    "不替代DCS/SIS；同时记录方案、确认、执行和复盘，使调度经验能够被检索和复用。首批范围聚焦产供销协同，"
    "稳定运行提示和关键设备趋势异常作为辅助能力。"
)

PART2 = (
    "我们建设“氨智领航调控师”，采用“事实汇总、约束校核、方案比较、人工确认、岗位接令、执行跟踪、班后复盘”流程。"
    "首先校核ERP订单导出、MES班次计划、罐区历史摘要和飞书协同记录四类离线或只读数据，形成带来源、时间戳和质量码的班次调度事实表。"
    "审批不由汇总分数驱动：一般偏差只提示，设备、公辅、库存等明确条件触发相关专业会签；命中演练停算线时停止新目标、经济排序、动作单和审批，维持最近批准方案并转现场规程。"
    "审批后由合成主操、硝酸主操、罐区/调度和公辅调度分别接令，全部接令后才进入执行跟踪，不写DCS/SIS。"
    "接令不等于实际执行，无Historian回传时实际值为空。班后记录执行偏差和未采纳原因，候选规则经专业人员审核后再入库。试点分三步实施：30天完成四类离线或只读数据的口径校核和接口清单，"
    "60天扩展生产历史数据库、设备和公用工程只读数据并开展旁路回放，90天选择条件完整且未命中停算线的场景进行执行验证。"
)

FEISHU_DOC_URL = "https://www.feishu.cn/docx/AYyad50itooOPxxZpQacgSqIn2c"
FEISHU_BASE_URL = "https://www.feishu.cn/base/QzENbAkl1aYQGds8dBacqu6Inue"
FEISHU_TASK_URL = "https://applink.feishu.cn/client/todo/task_list?guid=2ab6c357-dfeb-4f75-9aa3-781dc7ac7244"
FEISHU_VERIFIED_TASK_URL = "https://applink.feishu.cn/client/todo/detail?guid=44631b59-b834-47b1-a413-b751f2f291da&suite_entity_num=t136777"

FINAL_PRIORITY_LANES = [
    {
        "priority": "核心业务",
        "name": "产供销协同与跨装置负荷联动",
        "focus": "液氨有限、下游需求、市场价格、能源波动、装置负荷之间实时判断。",
        "answer": "给出液氨分配、装置负荷、外采安排和停运评估建议。",
        "demo": "液氨去向重分配及硝酸、尿素、纯碱、复合肥、外售的相对外售分配贡献排序。",
    },
    {
        "priority": "稳产辅助",
        "name": "跨装置波动协调",
        "focus": "下游需求突变、公用工程约束、能源窗口和库存波动对主流程的影响。",
        "answer": "给出跨装置负荷调整顺序，并列出保持连续生产所需的边界。",
        "demo": "下游降负荷后合成氨负荷、液氨库存和公用工程的联动节奏。",
    },
    {
        "priority": "设备辅助",
        "name": "关键设备趋势异常提示",
        "focus": "压缩机、气轮机驱动系统、关键机组和合成塔的持续趋势偏离。",
        "answer": "显示偏离变量、持续时间、核对点位、检查步骤和通知专业。",
        "demo": "振动、轴位移、防喘振裕度、轴承温度和床层温升的趋势证据。",
    },
    {
        "priority": "经验管理",
        "name": "经审核的调度案例与规则",
        "focus": "班组处置、未采纳原因和复盘结论缺少统一记录。",
        "answer": "按场景记录事实、判断、执行结果和适用条件，经专业人员审核后复用。",
        "demo": "飞书多维表格复盘库、智能问答和新人场景训练清单。",
    },
]

EXISTING_SYSTEM_STACK = [
    {
        "layer": "运行管理层",
        "system": "MES与合成氨调控平台",
        "site_note": "企业访谈中记录为“海尔豪斯”调控平台，具体厂商名称待企业确认。",
        "role": "承接日计划、班次执行、交接摘要、负荷调整记录和复盘闭环。",
        "ai_boundary": "审批通过后只写计划摘要、交接说明和复盘结果。",
    },
    {
        "layer": "数据采集层",
        "system": "IoT平台",
        "site_note": "采集合成氨DCS底层数据，形成可供上层系统读取的过程数据底座。",
        "role": "读取负荷、温度、压力、流量、电耗、蒸汽、罐区等实时/准实时数据。",
        "ai_boundary": "作为平台的只读事实源，不替代DCS控制。",
    },
    {
        "layer": "生产控制层",
        "system": "合成氨DCS",
        "site_note": "企业访谈中记录为中控体系，具体名称待企业确认。",
        "role": "完成合成氨主装置的局部回路控制、顺控、报警和联锁保护。",
        "ai_boundary": "通过生产历史数据库或只读接口获取摘要，不直接写控制参数。",
    },
    {
        "layer": "热电与公辅控制层",
        "system": "和利时DCS与热电APC",
        "site_note": "热电侧采用和利时DCS，并已有APC先进控制用于热电管网优化。",
        "role": "提供蒸汽、电力、公用工程负荷、管网约束和热电优化结果。",
        "ai_boundary": "读取热电APC约束余量和能耗窗口，生成跨装置调度建议。",
    },
    {
        "layer": "机组与压机层",
        "system": "SMC压机软件、康迪森机组数据",
        "site_note": "企业访谈中记录机组数据已接入和利时平台，具体点位清单待企业确认。",
        "role": "提供离心压缩机、气轮机驱动系统和关键机组的状态与趋势数据。",
        "ai_boundary": "用于设备趋势异常提示和停运评估，不直接触发开停车。",
    },
]


DELIVERY_STATUS_ROWS = [
    ["已验证", "本地交互工作台、飞书在线方案稿、多维表格样例库、Task v2任务清单和验收任务t136777。"],
    ["原型", "互动卡片、负荷调整审批状态机、典型班次证据链、事件回传字段契约和班后案例审核流程。"],
    ["待授权", "目标调度群机器人、审批定义、Aily知识源、事件订阅，以及ERP、MES、生产历史数据库、APC、设备和公用工程接口。"],
]


CONDITION_ROUTING_ROWS = [
    ["一般提示", "数据有效，且未命中专业会签条件或演练停算线", "显示具体偏差、持续时间和核对事项；不自动创建审批", "调度员"],
    ["相关专业会签", "设备趋势达到specialist_review条件，公辅余量异常或失效，库存接近确认边界，或其他已登记条件成立", "显示触发字段、时间窗、质量状态和规则版本，通知对应专业", "设备、公辅、罐区或工艺专业"],
    ["演练停算线", "趋势回放标记stop_new_recommendation_demo，或接口矩阵登记的停算条件成立", "停止新目标、经济排序、动作单和审批；维持最近批准方案并转现场规程", "现场指挥链与对应专业"],
]


DATA_SAMPLE_ROWS = [
    ["data/compressor_trend_replay_sample.csv", "压缩机趋势回放", "负荷、轴振、轴位移、防喘振裕度、轴承温度、质量码和复核标签", "仅验证观察、专业会签和演练停算流程；不作为企业设备阈值"],
    ["data/interface_field_matrix.csv", "接口字段矩阵", "阶段、来源系统、字段、方向、频率、时间基准、质量规则、单位、责任专业和超时动作", "字段与动作是定义稿，30天校核后由企业冻结"],
]


DEMO_SCRIPT_ROWS = [
    ["0:00-0:35", "总览与当班工作区", "先说明安全、连续生产和经济比较的顺序，明确平台位于生产运行管理层。", "系统边界、当班状态、数据质量状态"],
    ["0:35-1:15", "行情核价与班次事实", "说明未经经营确认的行情不进入新方案，每个输入保留来源、时间、单位和质量状态。", "行情版本、确认状态、证据完整度"],
    ["1:15-2:10", "24小时平衡与方案比较", "先核对液氨守恒和安全库存，再比较各去向；经济指标不表述为已实现收益。", "物料平衡、候选方案、逐项贡献账本"],
    ["2:10-2:55", "设备趋势处置", "展示具体变量、持续时间和检查项；达到specialist_review条件时通知设备专业会签。", "趋势回放样例、触发字段、检查清单"],
    ["2:55-3:50", "条件处置与飞书闭环", "一般提示不自动升级；明确条件触发专业会签；演练停算线停止新目标、经济排序、动作单和审批。", "处置状态、触发条件、会签对象、停算动作"],
    ["3:50-4:35", "执行跟踪与班后复盘", "四岗位全部接令后才进入执行跟踪；接令不等于执行，无Historian回传时实际值为空。", "动作单版本、回传状态、执行偏差、审核状态"],
    ["4:35-5:00", "试点与验收", "说明30/60/90天范围、有效班次定义和必过项，再说明效果观察指标。", "数据清单、验收口径、停用条件"],
]


EVIDENCE_CHAIN_ROWS = [
    ["E01", "原班次计划、触发时间和输入快照哈希", "调度员", "班次调度事实表", "原型"],
    ["E02", "订单、MES计划、罐区摘要、行情确认和数据质量状态", "经营、生产、罐区", "班次调度事实表", "样例已形成；企业数据待授权"],
    ["E03", "物料守恒、安全库存、公辅余量、设备趋势和条件路由结果", "生产、设备、安环", "约束校核记录", "原型"],
    ["E04", "原计划、候选方案、有效期、撤销条件和贡献账本", "调度员、财务", "方案版本与测算明细", "原型"],
    ["E05", "班长确认，以及命中明确条件时对应专业的会签或退回意见", "班长、调度主管、对应专业", "飞书审批记录", "状态机为原型；审批定义待授权"],
    ["E06", "动作单版本、四岗位接令人、接令时间、回复和异常", "四岗位主操/调度", "Task v2与动作单记录", "普通任务已验证；自动拆单为原型"],
    ["E07", "Historian回传的实际负荷、库存、能耗和执行偏差", "生产、信息化", "执行跟踪记录", "待授权；接令不作实际值"],
    ["E08", "未采纳原因、归因、适用条件、审核人和规则版本", "生产、设备、安环、经营", "Base待审核案例库", "样例库已验证；审核流程为原型"],
]


DISPATCH_ACTION_ROWS = [
    ["合成主操", "目标负荷、升降负荷节奏、保持条件和撤销条件", "确认收到有效版本，按现行操作票组织执行", "保留原计划并反馈不可执行原因"],
    ["硝酸主操", "下游目标负荷、耗氨窗口、缓冲能力和恢复条件", "确认下游承接节奏与现场条件", "退回调度重算液氨去向与负荷方案"],
    ["罐区/调度", "液氨库存、收发安排、安全库存和物流窗口", "确认库存与收发计划能够支撑动作单", "标记冲突并冻结执行跟踪入口"],
    ["公辅调度", "蒸汽、电力、空分、循环水等余量和时间窗口", "确认公辅条件满足动作单要求", "返回受限时段和可用窗口，重新排程"],
]


JUDGE_QA_ROWS = [
    ["与普通看板有什么区别？", "平台不只展示状态，还保存原计划、输入快照、约束校核、候选方案、人工确认、执行回传和班后复盘，使每条建议可追溯。"],
    ["为什么不直接控制DCS？", "DCS、SIS、APC和设备保护承担控制与保护。本平台处理班次级跨装置决策，审批后最多写MES计划摘要、交接说明和复盘记录。"],
    ["哪些飞书能力已经做通？", "已验证在线方案稿、Base样例库、Task v2任务清单和任务t136777；卡片、审批状态机和事件契约为原型；Aily、目标群和企业接口待授权。"],
    ["没有实时数据如何证明有效？", "当前验证业务流程、计算口径、控制边界和可追溯性。现场效果要经过数据校核、旁路回放和条件受控执行验证后计算。"],
    ["经济指标为什么不是节约金额？", "演示值用于比较液氨去向。只有建议被采纳、实际执行并完成归因后，才进入试点效果统计。"],
    ["如何防止错误建议？", "关键字段缺失、数据过期、口径冲突、超出适用工况或接近安全设备限值时，平台保留原计划并转人工复核。"],
    ["审批是否由汇总分数决定？", "不是。一般提示、专业会签和演练停算线分别由具体字段、时间窗、质量状态和规则版本决定；停算时不生成新审批。"],
    ["系统怎样学习而不越学越错？", "班后案例须经生产、设备、安环和经营专业审核后才可修订规则或纳入训练样本，单个班次不自动改变生产规则。"],
    ["与MES、APC和RTO是什么关系？", "MES承接计划和记录，APC处理装置内先进控制，RTO负责授权范围内的经济优化；本平台组织跨装置班次事实、方案比较和人工确认。"],
    ["90天怎样验收？", "先通过安全、数据真实性、可复算和审批合规等必过项，再覆盖不少于30个有效班次和三类场景，比较双方确认的观察指标。"],
    ["最先需要企业提供什么？", "首批只需订单导出、MES班次计划、罐区历史摘要和飞书协同记录，并由相关专业确认字段口径、时间戳、单位和更新频率。"],
]


PILOT_DATA_ROWS = [
    ["30天", "订单与经营", "产品、数量、交期、优先级、价格版本、确认人", "经营、销售", "离线导出校核，待授权"],
    ["30天", "MES班次计划", "班次、装置、原计划负荷、目标产量、计划版本、变更原因", "生产、调度", "离线导出校核，待授权"],
    ["30天", "罐区历史摘要", "期初量、收发量、期末量、安全库存、单位、质量状态", "罐区、计量", "历史摘要校核，待授权"],
    ["30天", "飞书协同记录", "建议编号、确认人、负责人、截止时间、状态、未采纳原因", "调度、班组", "Task v2和Base样例已验证；自动链路为原型"],
    ["60天", "生产历史数据库", "合成氨负荷、液氨产量与库存、下游耗氨量、关键温压流、质量码", "生产、信息化", "只读接入待授权"],
    ["60天", "公辅与能源", "蒸汽、电力、燃料、空分与管网余量、约束状态", "热电、公辅", "只读接入待授权"],
    ["60天", "关键设备", "振动、轴位移、防喘振裕度、轴承温度、机组状态、检修记录", "设备、仪表", "只读接入待授权"],
    ["60天", "财务测算", "非氨变动成本、包装物流、启停损失、机会成本、违约避免额", "财务、经营", "口径待企业确认"],
]


ACCEPTANCE_ROWS = [
    ["安全边界", "安全硬约束零突破；不存在DCS/SIS写入路径，不替代报警、联锁、操作票和现场确认。", "必过"],
    ["数据真实性", "关键字段缺失、过期或质量无效时转人工；无生产历史数据库回传时实际值为空。", "必过"],
    ["可复算性", "液氨物料守恒、相对外售分配贡献和方案排序可按同一输入快照复算。", "必过"],
    ["条件处置", "一般提示、专业会签和演练停算线可追到字段、时间窗、质量状态和规则版本；审批不由汇总分数驱动。", "必过"],
    ["动作单门禁", "审批后四岗位分别接令，未全部接令不得进入执行跟踪；接令状态不得填入实际执行字段。", "必过"],
    ["场景覆盖", "90天不少于30个有效班次，覆盖产供销重排、公辅波动吸收和设备趋势复核。", "必过"],
    ["效果观察", "按双方冻结口径比较数据完整率、调度用时、订单满足率、库存偏差、单位氨变动成本和提示提前量。", "基线确认后评价"],
]


ACCEPTANCE_TARGET_ROWS = [
    ["数据可用率", "新鲜度窗口内且质量状态合格的必需记录数 / 应到必需记录数", "不低于98%"],
    ["物料平衡残差", "|期初可用量+期间产量+外采-各去向-期末可用量| / (期初可用量+期间产量+外采)", "不高于1.5%"],
    ["快照重算P95", "合格输入快照冻结至方案、约束证据和测算明细生成完成的95分位耗时，不含人工确认", "不高于120秒"],
    ["调度决策中位耗时缩短", "1-试点同类场景决策中位耗时/基线中位耗时", "不低于30%"],
]


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = FONT_CN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def setup_doc(title: str, subtitle: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.86)
    section.right_margin = Inches(0.86)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(5)

    for name, size in [("Heading 1", 15), ("Heading 2", 12.5), ("Heading 3", 11.5)]:
        style = styles[name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
        style.font.size = Pt(size)
        style.font.color.rgb = BLUE
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.text = "氨智领航调控师 | 合成氨生产调度专家平台"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], size=8.5, color=GRAY)

    section.footer.paragraphs[0].text = ""

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run_font(r, size=21, bold=True, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_run_font(r, size=10.5, color=GRAY)
    p.paragraph_format.space_after = Pt(15)
    return doc


def para(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r)


def numbers(doc: Document, items: list[str]):
    numbering = doc.part.numbering_part.element
    style = doc.styles["List Number"]._element
    style_num_id = style.find(f".//{qn('w:numId')}")
    base_num_id = int(style_num_id.get(qn("w:val")))
    base_num = next(node for node in numbering.findall(qn("w:num")) if int(node.get(qn("w:numId"))) == base_num_id)
    abstract_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    new_num_id = max(int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    for item in items:
        p = doc.add_paragraph(style="List Number")
        num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), str(new_num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id)
        r = p.add_run(item)
        set_run_font(r)


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for idx, header in enumerate(headers):
        cell = tbl.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        cell.width = Inches(widths[idx])
        r = cell.paragraphs[0].add_run(header)
        set_run_font(r, bold=True)
    for row in rows:
        cells = tbl.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[idx].width = Inches(widths[idx])
            set_cell_margins(cells[idx])
            r = cells[idx].paragraphs[0].add_run(value)
            set_run_font(r, size=9.2)
    return tbl


def add_references(doc: Document, indices: list[int] | None = None):
    doc.add_heading("参考文献与资料来源", level=1)
    selected = REFERENCES if indices is None else [REFERENCES[i] for i in indices]
    for i, ref in enumerate(selected, 1):
        para(doc, f"[{i}] {ref}")


def save(doc: Document, filename: str):
    MATERIALS.mkdir(exist_ok=True)
    path = MATERIALS / filename
    try:
        doc.save(path)
    except PermissionError:
        path = path.with_name(f"{path.stem}_决赛更新稿{path.suffix}")
        doc.save(path)
        print(f"原文件正在使用，已另存：{path.name}")
    return path


def add_dispatch_action_order(doc: Document, heading: str, level: int = 1):
    doc.add_heading(heading, level=level)
    para(
        doc,
        "审批通过后生成跨装置调度动作单，由合成主操、硝酸主操、罐区/调度和公辅调度分别接令。四个岗位全部接令后才可进入执行跟踪；接令只表示收到有效方案并确认责任，不表示现场已经执行。",
    )
    table(
        doc,
        ["接令岗位", "接令时核对", "接令后的责任", "退回处理"],
        DISPATCH_ACTION_ROWS,
        [1.0, 2.1, 2.15, 1.15],
    )
    para(
        doc,
        "动作单以班次编号、方案版本和岗位标识唯一记录接令人、接令时间、回复、异常和有效期。审批状态机与动作单门禁为原型，Task v2普通任务写入已验证；自动拆单、岗位身份映射和事件汇总待授权。实际负荷、库存和能耗只接受Historian或企业认可的只读事实源回传，无回传时保持为空。",
    )


def build_opening_report():
    doc = setup_doc("开题报告", "命题：如何用AI打造合成氨生产调度专家")
    doc.add_heading("Part 1：命题前置分析与洞察", level=1)
    para(doc, PART1)
    doc.add_heading("Part 2：整体解决方案设计", level=1)
    para(doc, PART2)
    doc.add_heading("命题理解", level=1)
    table(
        doc,
        ["命题关注", "本方案回应", "可验证结果"],
        [
            ["调度寻优难", "用班次调度事实表统一订单、库存、能耗、设备和工艺约束，比较液氨分配与负荷调整方案。", "每班保留原计划、调度建议、确认结果和执行偏差。"],
            ["指挥效率低", "将调度建议转为飞书卡片、审批单、交接摘要和执行监督窗口。", "班前会只看差异、约束和确认事项。"],
            ["知识传承难", "记录采纳、驳回、异常处置和班长判断，专业人员审核后进入案例库。", "周度复盘保留规则修订建议、审核人和版本。"],
            ["试点实施", "先做离线与只读数据校核，再开展旁路回放和条件受控执行验证。", "30/60/90天范围、停算条件和预登记验收目标明确。"],
        ],
        [1.2, 2.75, 2.45],
    )
    add_references(doc, [1, 2, 6, 9, 13, 14])
    return save(doc, "01_开题报告.docx")


def build_solution_doc():
    doc = setup_doc("整体解决方案书", "氨智领航调控师：产供销全线联动数字调度长")
    doc.add_heading("1. 方案定位", level=1)
    para(
        doc,
        "氨智领航调控师部署在生产运行管理层，汇总液氨库存、下游订单、产品行情、公用工程和设备状态，生成班次级分配与负荷建议。DCS、SIS、APC和设备保护系统继续承担控制与保护，平台仅提供只读分析、人工审批和执行记录。",
    )
    doc.add_heading("2. 业务主线", level=1)
    table(
        doc,
        ["优先级", "方向", "重点问题", "演示抓手"],
        [[lane["priority"], lane["name"], lane["focus"], lane["demo"]] for lane in FINAL_PRIORITY_LANES],
        [0.9, 1.55, 2.55, 1.4],
    )
    doc.add_heading("3. 总体架构", level=1)
    table(
        doc,
        ["层级", "核心对象", "交付内容"],
        [
            ["班次数据", "订单导出、MES班次计划、罐区历史摘要、飞书协同记录；后续扩展生产历史数据库与设备数据", "班次调度事实表、口径字典、数据质量报告"],
            ["约束与候选模型", "机理硬约束、设备趋势、库存/订单预测；PINN仅作为待验证的反应器校准候选", "规则版本、适用范围、漂移监控、留出班次回放报告"],
            ["联动调度", "液氨机会成本、非氨变动成本、流程连续性、APC/MPC候选建议", "液氨分配、负荷调整、外采安排、停运评估和撤销条件"],
            ["协同执行", "飞书卡片、审批、多维表格、交接摘要", "确认人、审批人、执行时间、采纳/驳回原因"],
            ["经验管理", "经审核的调度案例、异常处置、未采纳原因和新人训练样本", "规则修订建议、审核记录、版本和适用条件"],
        ],
        [1.1, 2.35, 3.05],
    )
    doc.add_heading("4. 现有系统对接口径（待企业确认）", level=1)
    para(
        doc,
        "企业访谈记录显示，现场已有或正在建设多套生产与经营系统。平台在现有系统之间汇总班次调度所需信息，并明确数据来源、流向和写回边界；DCS、MES、APC和压机软件的原有职责保持不变。",
    )
    table(
        doc,
        ["层级", "现场系统", "主要作用", "平台边界"],
        [[item["layer"], item["system"], item["role"], item["ai_boundary"]] for item in EXISTING_SYSTEM_STACK],
        [1.05, 1.55, 2.55, 1.25],
    )
    bullets(doc, [f"{item['system']}：{item['site_note']}" for item in EXISTING_SYSTEM_STACK])
    doc.add_heading("5. 飞书功能模块在方案中的位置", level=1)
    para(
        doc,
        "飞书承接调度建议送达、审批、任务跟踪和班后复盘。调度建议完成责任人确认和状态留痕后，才进入MES计划摘要或复盘记录。",
    )
    table(
        doc,
        ["飞书功能", "解决的命题问题", "合成氨场景中的动作", "沉淀数据"],
        [
            ["互动卡片", "指挥效率低", "把目标负荷、处置状态、触发字段、相对外售分配贡献和确认按钮推送到当班调度群。", "card_action、operator_confirm、reject_reason"],
            ["审批", "条件会签与责任链", "设备、公辅、库存等登记条件成立时，路由班长、主管和对应专业确认。", "approval_instance、approval_status、trigger_condition"],
            ["多维表格Base", "知识传承难", "记录采纳/驳回、实际负荷、单位氨能耗、库存变化、未采纳原因和班后复盘。", "shift_fact、execution_delta、lesson_learned"],
            ["飞书任务", "执行监督弱", "企业发布后，由审批结果触发跟踪任务，核对生产历史数据库中的实际负荷、库存变化和班后结论。", "owner、due_time、checklist_status、close_note"],
            ["飞书智能问答/事件回传", "建议解释与案例复用", "授权后支持调度员追问方案依据；操作结果返回待审核案例库，不自动修改生产规则。", "query_log、callback_event、idempotency_key"],
        ],
        [1.0, 1.25, 2.55, 1.6],
    )
    doc.add_heading("6. 与APC/RTO体系的接口关系", level=1)
    bullets(
        doc,
        [
            "APC侧：平台不覆盖底层控制回路，只读取关键受控变量和约束余量，输出班次级目标负荷、升降负荷速率建议和需保持的边界条件。",
            "RTO侧：平台把天然气/煤气化、蒸汽、电力、液氨库存占用、订单延期、外售机会成本、边际贡献和固定成本吸收转化为经济目标，供实时优化或离线回放使用。",
            "机理与模型侧：先以合成塔热平衡、反应器温升、氢氮比、循环气量和压缩机效率形成硬约束；PINN完成企业历史回放与留出班次验证后，才可作为旁路计算的候选校准方法。",
            "MES侧：审批通过后只写入计划、交接摘要和复盘记录，不直接写入DCS/SIS控制参数。",
        ],
    )
    doc.add_heading("6（一）条件驱动处置", level=2)
    para(doc, "审批不由任何汇总分数驱动。平台逐项核对字段条件、持续时间、质量状态和规则版本，并按一般提示、相关专业会签和演练停算线三类处置。")
    table(doc, ["处置", "明确条件", "系统动作", "责任岗位"], CONDITION_ROUTING_ROWS, [1.0, 2.45, 2.15, 0.8])
    doc.add_heading("7. 典型场景", level=1)
    table(
        doc,
        ["场景", "触发条件", "建议动作", "试点观察项"],
        [
            ["高优订单保障", "尿素溶液或复合肥用氨急单、液氨库存偏低", "比较提高合成负荷、压缩弹性外售和调整其他下游分配。", "订单延期风险、液氨期末库存和相对外售分配贡献。"],
            ["能源错峰", "蒸汽、电力或原料边际成本变化", "比较不同负荷时段，保持企业确认的安全库存。", "单位氨变动成本、库存占用和负荷变化次数。"],
            ["设备趋势异常", "设备指标持续偏离或合成塔温升异常", "显示趋势证据，建议复核负荷上限和点检窗口。", "提示提前量、检查结果和人工处置记录。"],
            ["转人工处理", "数据过期、口径冲突或输入超出模型适用工况范围", "停止给出可执行建议，不回写计划，保留原计划并提示人工复核。", "降级原因、恢复条件和复核用时。"],
        ],
        [1.0, 2.0, 2.45, 1.45],
    )
    doc.add_heading("8. 跨系统班次信息汇总", level=1)
    para(
        doc,
        "企业访谈显示，班次调度信息分散在MES、DCS、IoT、APC、压机软件、热电平台和经营数据中。调度长需要逐项查看价格、库存、订单、装置负荷、设备状态和公用工程约束。平台按班次汇总这些信息，并记录每次判断使用的数据和规则。",
    )
    table(
        doc,
        ["现场问题", "当前做法", "平台处理", "试点观察项"],
        [
            ["信息分散", "从多套系统逐项查看。", "生成带来源、时间戳和质量码的班次调度事实表。", "查数用时和口径差异数量。"],
            ["判断依赖经验", "由调度长判断哪些信息影响当班方案。", "记录历史方案、驳回原因、复盘结论和适用条件。", "新人场景判断准确率与用时。"],
            ["跨线取舍复杂", "人工比较液氨去向、下游负荷、热电约束和设备状态。", "列出候选分配方案、约束、牺牲项和撤销条件。", "人工复核差异和决策用时。"],
            ["复盘记录不完整", "判断常停留在口头交接。", "通过飞书任务、审批和多维表格保存待审核案例。", "复盘字段完整率和审核完成率。"],
        ],
        [1.15, 1.8, 2.45, 1.0],
    )
    doc.add_heading("9. 产供销经济取舍口径", level=1)
    para(
        doc,
        "产供销调度采用增量口径比较液氨去向：产品净收入减非氨变动成本、包装物流、液氨机会成本和启停摊销，再加经财务确认的违约避免额。结果称为“相对外售分配贡献测算值”，用于方案比较，不表示已经实现的收益。折旧和固定人工仅作管理解释，不重复计入。",
    )
    table(
        doc,
        ["测算项", "现场含义", "调控师使用方式"],
        [
            ["销售价", "尿素、纯碱、硝酸、复合肥、液氨外售等实时或预测行情。", "识别价格趋势，但不单独作为开停依据。"],
            ["变动成本", "原料、蒸汽、电力、包装、物流等随产量变化的成本。", "判断产品是否至少覆盖直接增量成本。"],
            ["边际贡献", "销售收入扣除变动成本后对折旧、人工、公辅和固定费用的贡献。", "售价低于完全成本但边际贡献为正时，可作为“继续开”的理由之一。"],
            ["固定成本吸收", "连续开车可能覆盖部分折旧、固定人工和装置维持费用。", "作为管理层解释单列展示，不再计入增量贡献，避免重复核算。"],
            ["机会成本", "液氨有限时，供硝酸、尿素、纯碱、复合肥或外售之间的替代收益。", "选择单位液氨综合贡献更高且不破坏安全连续性的去向。"],
            ["停开成本", "停车、重启、废料、人员和物料损失。", "当边际贡献接近零时，停开成本可能决定是否继续运行。"],
        ],
        [1.05, 1.95, 3.4],
    )
    doc.add_heading("10. 跨装置联动的调度优先级", level=1)
    para(
        doc,
        "企业访谈确认：合成氨调度不能只看当班利润，必须先保安全，再保流程连续，只有在流程确实需要中断时才进入特殊取舍。DCS可以完成液位、阀门、局部顺控和联锁保护，但难以自动判断下游异常、气区设备、能源窗口和产品结构变化后的跨大单元联动，因此本方案把这部分做成调控师的规则层。",
    )
    table(
        doc,
        ["优先级", "判断口径", "调控师动作", "留痕与边界"],
        [
            ["1. 安全", "安环红线、SIS/DCS联锁，或罐区、压缩机、合成塔关键测量接近企业确认边界。", "命中演练停算线时停止新目标、经济排序、动作单和审批，维持最近批准方案。", "不绕过DCS/SIS；转现场规程并记录触发字段。"],
            ["2. 流程不中断", "气化、净化、合成、空分、气轮机驱动系统和液氨去向的连续性。", "优先做跨装置负荷联动，减少停车后重启时间、废料、人员和物料损失。", "输出重启成本、废料成本、影响范围和回滚条件。"],
            ["3. 特殊取舍", "流程确需中断时，比较空分、合成气轮机驱动系统与下游装置的停运代价。", "按安全、重启复杂度、废料损失和主流程影响排序；企业口述案例仅作为待复核候选，不固化为通用停机规则。", "由班长、调度主管和相关专业会签；平台不自动执行停机。"],
        ],
        [0.9, 2.0, 2.35, 1.25],
    )
    doc.add_heading("11. 30/60/90天落地路径", level=1)
    table(
        doc,
        ["阶段", "目标", "交付物", "验收口径"],
        [
            ["30天", "校核四类离线或只读数据", "订单导出、MES班次计划、罐区历史摘要、飞书协同记录的字段字典、样例和质量报告", "逐字段确认来源、时间戳、单位、质量状态和责任专业，不接入DCS实时写路径。"],
            ["60天", "扩展只读数据并开展旁路回放", "生产历史数据库、设备和公用工程只读接口；调度建议、人工方案及差异报告", "覆盖正常班次、急单、库存偏低和设备趋势异常场景。"],
            ["90天", "开展条件受控执行验证", "不少于30个有效班次，覆盖产供销重排、公辅波动吸收、设备趋势复核三类场景", "只选择未命中演练停算线且审批条件完整的场景；无Historian回传时实际值为空。"],
        ],
        [0.75, 1.65, 2.15, 2.05],
    )
    doc.add_heading("12. 验收与停用条件", level=1)
    bullets(
        doc,
        [
            "验收条件：双方先确认基线期、评价期和阈值；至少覆盖30个有效班次和产供销重排、公辅波动吸收、设备趋势复核三类场景。安全、数据质量、建议可复算和审批合规为必过项，再比较订单满足率、单位氨变动成本、库存占用、调度用时和提示提前量。",
            "贡献口径：24小时指标统一称为相对外售分配贡献测算值；实际效果只统计已采纳、已执行且完成归因的方案。",
            "停用条件：关键数据过期或缺失、输入超出适用工况范围、模型偏差超过双方确认的控制限，或安全与设备限值接近时停止给出可执行建议。",
        ],
    )
    doc.add_heading("12（一）预登记验收目标", level=2)
    table(doc, ["指标", "定义稿口径", "目标"], ACCEPTANCE_TARGET_ROWS, [1.35, 4.05, 1.0])
    para(doc, "以上目标是试点定义稿，不表示当前已达成。字段范围、分母、时间起止点、场景分层、基线期和排除项在30天数据校核后由企业冻结。")
    doc.add_heading("12（二）仓库内验收样例", level=2)
    table(doc, ["文件", "用途", "关键字段", "使用边界"], DATA_SAMPLE_ROWS, [1.65, 1.15, 2.35, 1.25])
    doc.add_heading("13. 典型班次证据链", level=1)
    table(
        doc,
        ["编号", "证据", "责任岗位", "保存位置", "状态"],
        EVIDENCE_CHAIN_ROWS,
        [0.55, 2.4, 1.25, 1.25, 0.95],
    )
    para(doc, "E01至E04完整后才允许提交建议；E05完成所需会签后生成动作单；E06四岗位未全部接令时不得进入执行跟踪；E07缺失时不得填写实际效果；E08未经审核不得修订规则。")
    add_dispatch_action_order(doc, "14. 跨装置调度动作单")
    doc.add_heading("15. 试点验收口径", level=1)
    table(doc, ["验收项", "口径", "判定"], ACCEPTANCE_ROWS, [1.05, 4.35, 1.0])
    doc.add_heading("Market source and execution-price gate", level=2)
    table(
        doc,
        ["Source", "Use", "Execution authority"],
        [
            ["NBS production materials", "Public benchmark; ten-day publication", "No"],
            ["CZCE UR/SA", "Authorized futures trend reference", "No"],
            ["MOFCOM commodity price", "Periodic industry cross-check", "No"],
            ["Yuntu ERP / settlement", "Product, region, tax and freight aligned business price", "Yes after confirmation"],
        ],
        [2.0, 3.2, 1.2],
    )
    para(doc, "Public reference refresh only updates evidence cards. An unconfirmed or stale price cannot enter a new executable plan or economic ranking; the last valid enterprise version is retained and the business owner is notified.")
    add_references(doc, [1, 2, 4, 5, 6, 9, 10, 11, 12, 13, 14])
    return save(doc, "02_整体解决方案书.docx")


def build_manual_doc():
    doc = setup_doc("调控师操作手册", "班次调度建议、飞书协同与转人工处理")
    doc.add_heading("1. 当班操作流程", level=1)
    table(
        doc,
        ["时间点", "调控师动作", "现场确认", "留痕位置"],
        [
            ["班前30分钟", "刷新订单、液氨库存、下游需求、行情、公用工程和设备状态，生成班次调度事实表。", "调度员核对来源、时间戳、单位和质量状态。", "班次调度事实表"],
            ["班前会", "查看产供销建议、跨装置协调事项和设备趋势异常提示。", "班长确认液氨分配、负荷调整、外采安排和撤销条件。", "飞书卡片"],
            ["班中偏差", "库存、设备、订单或能源条件变化时重新计算候选方案。", "未经确认继续执行原计划。", "方案变更记录"],
            ["班后复盘", "记录采纳状态、实际负荷、能耗、库存、偏差和原因。", "班长补充现场判断，专业人员审核后入库。", "飞书多维表格"],
        ],
        [1.0, 2.2, 2.0, 1.3],
    )
    doc.add_heading("1（一）当班使用顺序", level=2)
    numbers(
        doc,
        [
            "核对产供销数据：确认订单、库存、价格和相对外售分配贡献测算口径。",
            "检查连续生产约束：确认下游需求、公用工程、设备和库存变化是否需要跨装置协调。",
            "检查设备趋势：查看偏离变量、持续时间、数据来源、检查项和通知专业，不直接触发开停车。",
            "完成班后记录：保存采纳、驳回、执行偏差和现场原因，待专业人员审核后进入案例库。",
        ],
    )
    doc.add_heading("2. 飞书协同动作", level=1)
    bullets(
        doc,
        [
            "互动卡片原型：显示目标负荷、处置状态、触发字段、相对外售分配贡献、撤销条件和确认按钮。",
            "审批原型：设备、公辅、库存等明确条件成立时生成会签草稿，字段包括装置、目标负荷、触发字段、现场影响和撤销条件。",
            "飞书多维表格：记录班次事实、采纳状态、未采纳原因、执行偏差和复盘结论。",
            "飞书智能问答待授权：接入知识源后，可追问降负荷依据、库存可支撑时间和设备趋势检查步骤。",
        ],
    )
    table(
        doc,
        ["操作节点", "飞书动作", "班组确认口径"],
        [
            ["方案生成", "企业发布后，调控师向合成氨当班群发送互动卡片。", "核对目标负荷、触发字段、处置状态、撤销条件和与原计划差异。"],
            ["条件会签", "一般提示不自动升级；设备、公辅、库存等登记条件成立时增加主管和对应专业。", "未完成所需会签前不写MES计划摘要。"],
            ["执行监督", "企业发布后，审批结果触发飞书任务并列出核对清单。", "通过生产历史数据库只读数据核对实际负荷、库存和公用工程变化。"],
            ["班后记录", "复盘字段写入飞书多维表格，未采纳原因进入待审核案例库。", "记录不采纳原因和执行偏差，由专业人员决定是否修订规则。"],
        ],
        [1.1, 2.35, 2.95],
    )
    doc.add_heading("2（一）条件驱动处置", level=2)
    table(doc, ["处置", "明确条件", "系统动作", "责任岗位"], CONDITION_ROUTING_ROWS, [1.0, 2.45, 2.15, 0.8])
    para(doc, "命中演练停算线时，立即停止新目标、经济排序、动作单和审批，维持最近批准方案并转现场规程；解除条件由对应专业确认后再冻结新快照。")
    doc.add_heading("3. 异常处理原则", level=1)
    numbers(
        doc,
        [
            "安全与设备限值接近时，停止给出经济性负荷建议，保留原计划并提示人工处置。",
            "下游异常、公用工程约束或能源条件变化时，先核对气化、净化、合成和液氨去向的连续性。",
            "确需停运时，比较候选装置的重启时间、废料、人员和物料损失；既往案例只作提示，最终由相关人员会签。",
            "无法说明触发数据、约束、贡献测算和撤销条件的建议不得提交审批。",
            "任何负荷调整均由班组按现行操作票和审批制度执行；平台不写DCS/SIS，也不自动开停车。",
            "班后记录先进入待审核案例库，审核通过后才可修订规则或纳入训练样本。",
        ],
    )
    add_dispatch_action_order(doc, "4. 审批后的跨装置动作单")
    doc.add_heading("5. 班次留痕顺序", level=1)
    numbers(
        doc,
        [
            "建议生成前冻结原计划和输入快照，记录数据来源、时间、单位和质量状态。",
            "审批完成后生成四岗位动作单，任何岗位未接令或退回时保留原计划并返回调度重算。",
            "全部接令后进入执行跟踪，但不得把接令状态、任务完成按钮或人工备注填作实际负荷。",
            "Historian回传后计算建议值与实际值偏差；无回传时实际值为空，不进入效果统计。",
            "班后记录事实、决定、执行、原因和适用条件，经专业人员审核后归档。",
        ],
    )
    return save(doc, "03_调控师操作手册.docx")


def build_reference_doc():
    doc = setup_doc("参考文献与数据依据", "合成氨调度专家平台证据链")
    doc.add_heading("资料如何进入方案", level=1)
    para(doc, "本文件记录业务判断的来源、使用位置和适用边界。公司公告用于确认项目规模和产业链关系；行业标准用于划分系统边界和报警职责；工业软件资料用于对照功能，不作为本平台效果证明；控制优化研究用于支持方法选择。现场点位、阈值、成本和实际效果仍须由企业数据确认。")
    table(
        doc,
        ["资料类型", "支撑问题", "方案体现"],
        [
            ["云图公告与年报", "70万吨合成氨与下游肥化产业链协同", "以尿素溶液、复合肥、联碱、液氨外售作为调度对象。"],
            ["合成氨安全资料", "高温高压连续流程的安全边界", "DCS/SIS隔离、安环红线、罐区安全库存、停用条件。"],
            ["MPC/RTO与负荷调节研究", "连续装置动态优化可行性", "班次级目标负荷、经济目标、回放验证和模型治理。"],
            ["ISA-95与工业软件资料", "企业计划、制造运营与控制系统的职责边界", "生产管理层决策支持、只读取数和MES计划摘要边界。"],
            ["ISA-18报警管理资料", "区分DCS报警与平台一般提示", "设备趋势异常只作分析提示，不冒充控制系统报警。"],
            ["飞书Task v2与Aily资料", "任务跟踪和授权知识问答", "已验证任务API；智能问答、卡片、审批和事件回传按交付状态说明。"],
        ],
        [1.35, 2.2, 2.95],
    )
    doc.add_heading("关键事实提炼", level=1)
    bullets(
        doc,
        [
            "合成氨既是高能耗基础化工环节，也是氮肥产业链的关键原料入口，软件层运行优化具有持续价值。",
            "云图公告和2025年年度报告说明70万吨合成氨项目与下游肥化产业链存在协同关系，具体生产效果以企业数据为准。",
            "动态优化必须服从机理、安全和设备约束，不能把算法结果直接等同于可执行控制指令。",
            "企业试点从离线和只读数据校核开始，经过旁路回放、人工确认和停算条件验证后，再进入条件受控执行验证。",
        ],
    )
    doc.add_heading("仓库内验收样例", level=1)
    table(doc, ["文件", "用途", "关键字段", "使用边界"], DATA_SAMPLE_ROWS, [1.65, 1.15, 2.35, 1.25])
    para(doc, "试点预登记目标为数据可用率不低于98%、物料平衡残差不高于1.5%、快照重算P95不高于120秒、调度决策中位耗时缩短不低于30%。这些目标是定义稿，30天校核后由企业冻结。")
    add_references(doc)
    return save(doc, "04_参考文献与数据依据.docx")


def build_innovation_doc():
    doc = setup_doc("方案创新与落地清单", "从业务问题到试点交付")
    doc.add_heading("1. 已形成的交付基础", level=1)
    table(
        doc,
        ["能力", "当前状态", "企业试点工作"],
        [
            ["本地平台", "已完成交互工作台、24小时液氨平衡样例、相对外售分配贡献账本和流程演示。", "替换样例数据，确认现场点位、单位、阈值和财务口径。"],
            ["飞书协同", "在线文档、Base样例库和Task v2验收任务t136777已验证；卡片、审批和事件回传为原型。", "授权目标群机器人、审批定义、知识源、回调地址和企业应用权限。"],
            ["实施与治理", "已形成数据契约、控制边界、30/60/90天范围、停算条件和复盘字段。", "完成四类首批数据校核、旁路回放、条件受控执行验证和验收基线确认。"],
        ],
        [1.2, 2.35, 2.55],
    )
    doc.add_heading("2. 关键创新点", level=1)
    table(
        doc,
        ["方案特点", "解决的问题", "可验证方式"],
        [
            ["产供销协同", "把液氨库存、下游需求、市场价格、公用工程和装置负荷放在同一班次口径中比较。", "班次调度事实表、相对外售分配贡献排序和跨装置协调回放。"],
            ["跨装置波动协调", "下游需求或公用工程条件变化时，给出负荷调整顺序和连续生产边界。", "需求变化回放、负荷节奏建议和撤销条件。"],
            ["机理约束与候选校准", "先把反应器温升、氢氮比、循环气量和压缩机效率作为硬边界；PINN不预设为已交付能力。", "企业历史回放、留出班次验证、适用域和退出条件。"],
            ["APC/RTO接口边界", "读取约束余量和候选经济目标，不改写底层控制器。", "接口清单、变量映射、写回边界和审批记录。"],
            ["飞书协同", "区分已验证链路、原型和待授权能力。", "Task v2任务、Base样例、卡片与审批原型、Aily和事件回传授权清单。"],
            ["未采纳原因记录", "把专家否决和现场判断保存为待审核案例。", "每周输出候选规则、审核人、适用条件和版本。"],
            ["与原计划对照的贡献核算", "避免把行情变化误写为平台带来的效果。", "保留原计划、调度建议、确认结果、实际执行和外部影响。"],
        ],
        [1.4, 2.45, 2.35],
    )
    doc.add_heading("3. 首批上线清单", level=1)
    bullets(
        doc,
        [
            "30天数据：订单导出、MES班次计划、罐区历史摘要和飞书协同记录的字段字典、样例与质量报告。",
            "60天验证：生产历史数据库、设备和公用工程只读数据的旁路回放、人工方案对照和偏差报告。",
            "90天执行：经授权的飞书审批与任务跟踪、MES计划摘要、班后复盘和停用条件验证。",
            "治理：DCS/SIS无写入路径、人工确认、版本审核和相对外售分配贡献核算。",
        ],
    )
    add_references(doc, [1, 2, 9, 10, 11, 12, 13, 14])
    return save(doc, "05_方案创新与落地清单.docx")


def build_final_doc():
    doc = setup_doc("决赛完整方案文档", "云图控股合成氨命题 | 产供销全线联动数字调度长")
    doc.add_heading("一、参赛方案信息卡", level=1)
    table(
        doc,
        ["项目", "填写内容"],
        [
            ["队名", "氨智领航队"],
            ["命题", "如果来到全球领先的合成氨生产装置现场，如何用AI打造一个生产调度专家？"],
            ["方案标题", "氨智领航调控师：产供销全线联动数字调度长"],
            ["一句话摘要", "汇总订单、库存、生产历史、公用工程、设备和经营数据，在安全与连续生产约束下生成液氨分配、装置负荷、外采安排和停运评估建议，并通过飞书完成确认、跟踪和复盘。"],
            ["成员介绍与分工", "邓植斤：负责命题洞察、行业资料研读、系统架构、平台原型、飞书协同设计、提交材料与GitHub交付。"],
            ["使用的飞书AI能力", "已验证飞书Task v2与多维表格样例；互动卡片、审批和跨装置动作单门禁为原型；飞书智能问答、岗位身份映射与事件回传待授权。"],
        ],
        [1.45, 4.95],
    )
    doc.add_heading("二、方案成果展示", level=1)
    para(doc, "本方案已形成可本地打开的调控师工作台、样例数据、架构文档、飞书协同说明、操作手册和参考资料。工作台首先处理产供销协同，再辅助跨装置波动协调和设备趋势异常核对，班后记录经审核后进入案例库。")
    doc.add_heading("三、命题场景与痛点", level=1)
    bullets(doc, [
        "产供销平衡难：液氨有限、下游需求、市场行情、能源波动和装置负荷同时变化，调度长需要跨系统人工汇总后再判断。",
        "跨装置协调难：下游需求、公用工程或能源条件变化时，DCS负责局部回路和联锁，班组仍需判断不同装置的负荷调整顺序。",
        "经验传承难：成熟调度长培养周期长，专家判断常沉淀在经验和口头交接中，未采纳原因没有被结构化利用。",
    ])
    doc.add_heading("三（一）业务优先级", level=2)
    table(
        doc,
        ["优先级", "方向", "解决的问题", "Demo抓手"],
        [[lane["priority"], lane["name"], lane["answer"], lane["demo"]] for lane in FINAL_PRIORITY_LANES],
        [0.9, 1.55, 2.55, 1.4],
    )
    doc.add_heading("四、方案设计与差异", level=1)
    table(
        doc,
        ["模块", "能力", "不同于常规方案之处"],
        [
            ["产供销协同", "液氨去向、下游需求、价格、公用工程和装置负荷", "用统一口径比较候选分配方案。"],
            ["跨装置协调", "需求、能源、公用工程和库存变化", "给出负荷调整顺序和连续生产边界。"],
            ["约束与候选模型", "机理硬约束、APC/RTO接口、MPC建议；PINN仅作待验证校准候选", "先验证适用工况范围和退出条件，再进入旁路计算。"],
            ["班次执行", "调度建议、班次调度事实表和撤销条件", "将建议带入班前会、审批和MES计划摘要。"],
            ["跨系统汇总", "MES、IoT、生产历史数据库、APC、压机、热电和经营数据", "记录每个判断使用的数据来源、时间和质量状态。"],
            ["飞书协同", "卡片、审批、多维表格、Task v2、智能问答和事件回传", "按已验证、原型和待授权状态交付。"],
            ["案例审核", "采纳、驳回、偏差和外部影响记录", "专业人员审核后再修订规则或纳入训练样本。"],
        ],
        [1.1, 2.35, 2.95],
    )
    doc.add_heading("五、业务范围与落地价值", level=1)
    para(doc, "首批覆盖合成氨、液氨库存、尿素溶液、复合肥、联碱配套和液氨外售的协同调度。试点观察高优订单满足率、单位氨变动成本、液氨库存占用、调度用时、设备趋势提示提前量，以及相对外售分配贡献测算值与实际结果的偏差。")
    doc.add_heading("五（一）全产业链贯通价值", level=2)
    para(
        doc,
        "班次调度信息分散在MES、IoT、生产历史数据库、APC、压机软件、热电平台和经营数据中。调度长需要逐项核对液氨去向、下游负荷、能源条件、设备状态和订单优先级。平台将这些信息汇入班次调度事实表，并记录判断依据、确认人和执行结果。",
    )
    table(
        doc,
        ["现场问题", "平台处理", "试点观察项"],
        [
            ["信息分散", "汇总MES计划、生产历史数据、热电APC、公用工程、设备和市场信息。", "查数用时和口径差异数量。"],
            ["经验记录不完整", "记录采纳、驳回、复盘结论和适用条件，授权后供飞书智能问答检索。", "复盘完整率、新人场景判断准确率和用时。"],
            ["跨线取舍复杂", "在同一页面比较安全、连续生产、相对外售分配贡献和启停损失。", "人工复核差异和决策用时。"],
        ],
        [1.1, 3.0, 2.3],
    )
    doc.add_heading("五（二）边际贡献决策口径", level=2)
    para(
        doc,
        "平台采用可复算的增量口径比较液氨去向：产品净收入减非氨变动成本、包装物流、液氨机会成本和启停摊销，再加经财务确认的违约避免额。结果称为相对外售分配贡献测算值，不表示已实现收益；折旧和固定人工仅作管理解释。",
    )
    table(
        doc,
        ["判断项", "为什么重要", "平台输出"],
        [
            ["边际贡献", "判断该产品是否至少覆盖原料、蒸汽、电力、包装、物流等增量成本。", "继续开、降负荷或转供其他产品的经济理由。"],
            ["固定成本吸收", "连续生产可能覆盖折旧、固定人工和装置维持费用。", "单列解释，不计入增量贡献；是否继续生产仍由安全、连续与机会成本共同决定。"],
            ["液氨机会成本", "液氨有限时，供尿素、纯碱、硝酸、复合肥或外售的替代贡献不同。", "相对外售分配贡献排序和产品结构建议。"],
            ["停开成本", "停车重启会产生废料、时间、人员和物料损失。", "当边际贡献接近零时，给出继续运行或停产的取舍依据。"],
        ],
        [1.15, 2.75, 2.5],
    )
    doc.add_heading("五（三）跨装置调度的三条底线", level=2)
    table(
        doc,
        ["底线", "现场含义", "平台处理方式"],
        [
            ["安全第一", "任何经营目标都不能突破安环限值、DCS/SIS联锁边界和关键设备保护。", "数据不完整、输入超出适用工况范围或限值接近时，停止给出可执行建议并转人工处理。"],
            ["流程不中断第二", "连续流程一旦中断，重启耗时长、废料多、人员和物料损失高，尤其涉及气区和气轮机驱动系统。", "识别上下游异常后先做跨装置负荷联动，评估是否能通过液氨分配、下游降负荷或能源错峰吸收波动。"],
            ["特殊情况再取舍", "确需中断时，不是简单停哪套装置，而是比较空分、合成气轮机驱动系统和下游装置的总损失。", "把重启时间、废料和影响范围纳入排序；口述案例不固化为通用规则，最终由班长、调度主管和相关专业会签。"],
        ],
        [1.15, 2.55, 2.7],
    )
    doc.add_heading("五（四）现有系统接入假设", level=2)
    table(
        doc,
        ["系统", "在现场承担的作用", "本方案调用方式"],
        [
            ["MES/合成氨调控平台", "日计划、班次执行、负荷调整记录和交接复盘。", "审批通过后写计划摘要、交接说明和复盘结果。"],
            ["IoT平台", "采集合成氨DCS底层过程数据。", "作为班次调度事实表的只读数据源。"],
            ["合成氨DCS", "合成氨主装置局部回路控制、顺控、报警和联锁。", "通过生产历史数据库或只读接口获取摘要，不直接写控制参数。"],
            ["和利时DCS/热电APC", "热电、公用工程和蒸汽管网控制与优化。", "读取约束余量、能耗条件和APC结果，用于跨装置协调。"],
            ["SMC压机软件/康迪森机组数据", "压缩机、气轮机驱动系统和关键机组状态监测。", "用于设备趋势异常提示和停运评估。"],
        ],
        [1.55, 2.35, 2.5],
    )
    para(doc, "上述系统名称来自企业访谈记录，厂商及点位清单在企业试点前需由信息化、生产、设备和热电专业共同确认。")
    doc.add_heading("六、飞书功能接入", level=1)
    table(
        doc,
        ["飞书能力", "业务动作", "需要授权"],
        [
            ["互动卡片", "推送目标负荷、处置状态、触发字段和确认按钮。", "机器人发消息、目标群可见权限。"],
            ["审批", "设备、公辅、库存等登记条件成立时创建专业会签实例。", "approval_code、审批API权限、审批人范围。"],
            ["多维表格", "写入班次事实、采纳结果、未采纳原因、复盘结论。", "目标表格协作者权限和记录读写权限。"],
            ["飞书智能问答（Aily）", "调度员追问方案依据和异常处置步骤。", "知识源、问答权限和人工确认约束。"],
            ["事件回传", "接收卡片点击、审批通过或驳回、复盘提交。", "事件订阅、回调地址和签名校验。"],
        ],
        [1.2, 2.45, 2.75],
    )
    para(doc, f"飞书交付状态分为三类。已验证：在线方案稿 {FEISHU_DOC_URL}、多维表格样例库 {FEISHU_BASE_URL}、Task v2任务清单 {FEISHU_TASK_URL} 和验收任务 {FEISHU_VERIFIED_TASK_URL}。原型：互动卡片、审批状态机、跨装置动作单门禁和事件回传字段契约。待授权：目标群机器人、审批定义、四岗位身份映射、Aily知识源、事件订阅及MES和生产历史数据库接口。")
    doc.add_heading("六（一）飞书模块对命题三大挑战的支撑", level=2)
    table(
        doc,
        ["命题挑战", "飞书模块支撑", "形成的企业资产"],
        [
            ["调度寻优难", "授权后，飞书智能问答检索班次事实、历史方案和未采纳原因；事件回传把执行结果送入待审核案例库。", "可检索的调度案例、异常处置和候选校准样本。"],
            ["指挥效率低", "互动卡片把方案送到当班群，明确条件触发对应专业会签，任务监督跟踪执行与复盘。", "确认人、会签人、执行人、截止时间和关闭结果的责任链。"],
            ["知识传承难", "多维表格记录班次事实、未采纳原因、班长备注、执行偏差和外部影响；Aily授权后从审核知识源回答。", "带来源、审核人、适用条件和版本的班次案例。"],
        ],
        [1.15, 2.75, 2.5],
    )
    doc.add_heading("七、落地验证", level=1)
    numbers(doc, [
        "30天：校核订单导出、MES班次计划、罐区历史摘要和飞书协同记录四类离线或只读数据。",
        "60天：扩展生产历史数据库、设备和公用工程只读数据，开展旁路回放和人工方案对照。",
        "90天：覆盖不少于30个有效班次及产供销重排、公辅波动吸收、设备趋势复核三类场景；核对硬约束、可复算性、执行偏差和停用条件。",
    ])
    doc.add_heading("七（一）条件处置与预登记目标", level=2)
    table(doc, ["处置", "明确条件", "系统动作", "责任岗位"], CONDITION_ROUTING_ROWS, [1.0, 2.45, 2.15, 0.8])
    table(doc, ["指标", "定义稿口径", "目标"], ACCEPTANCE_TARGET_ROWS, [1.35, 4.05, 1.0])
    para(doc, "四项指标是试点定义稿，不表示当前已达成；30天数据校核后由企业冻结字段范围、分母、时间起止点、基线期和排除项。")
    doc.add_heading("八、方案体验入口与演示说明", level=1)
    para(doc, "体验入口：本地打开 D:\\云图-合成氨-邓植斤\\index.html，或访问 GitHub 仓库 https://github.com/zhijinDeng/synthesis-ammonia 查看完整静态平台、数据模型和提交材料。")
    para(doc, "演示中的价格、库存、负荷和设备趋势均为场景样例，不代表企业实时数据。现场数据、阈值、责任人和财务口径需在企业授权后确认。")
    doc.add_heading("八（一）五分钟演示脚本", level=2)
    table(doc, ["时间", "页面动作", "讲述要点", "当场证据"], DEMO_SCRIPT_ROWS, [0.8, 1.25, 2.85, 1.5])
    doc.add_heading("八（二）典型班次证据链", level=2)
    table(doc, ["编号", "证据", "责任岗位", "保存位置", "状态"], EVIDENCE_CHAIN_ROWS, [0.55, 2.4, 1.25, 1.25, 0.95])
    para(doc, "典型场景用于演示液氨库存趋紧、下游需求分化与设备趋势偏离时的跨系统判断，不作为云图现场工况结论。E06四岗位未全部接令时不得进入执行跟踪，E07缺失时不得填写实际效果。")
    add_dispatch_action_order(doc, "八（三）跨装置调度动作单", level=2)
    doc.add_heading("八（四）评委追问及事实回答", level=2)
    table(doc, ["追问", "事实回答"], JUDGE_QA_ROWS, [1.8, 4.6])
    doc.add_heading("八（五）试点数据与验收", level=2)
    table(doc, ["阶段", "数据域", "最小字段", "责任专业", "状态"], PILOT_DATA_ROWS, [0.65, 1.1, 2.8, 0.9, 0.95])
    table(doc, ["仓库样例", "用途", "关键字段", "使用边界"], DATA_SAMPLE_ROWS, [1.65, 1.15, 2.35, 1.25])
    para(doc, "有效班次须同时具备冻结的原计划、合格输入、可追溯建议、完整人工确认、Historian实际回传和班后复盘。未执行方案可用于流程验收，但不进入效果统计。")
    table(doc, ["验收项", "口径", "判定"], ACCEPTANCE_ROWS, [1.05, 4.35, 1.0])
    table(doc, ["指标", "定义稿口径", "目标"], ACCEPTANCE_TARGET_ROWS, [1.35, 4.05, 1.0])
    doc.add_heading("九、自由展示区", level=1)
    bullets(doc, [
        "项目已形成可交互页面、可复现生成脚本、数据模型、Word提交材料和GitHub开放仓库。",
        "材料使用班组可核对的业务对象、数据来源和操作边界描述平台能力。",
        "交付状态按已验证、原型和待授权三类标识，生产效果以试点数据为准。",
    ])
    doc.add_heading("十、附录", level=1)
    bullets(doc, [
        "附录A：参考文献与数据依据见 04_参考文献与数据依据.docx。",
        "附录B：调控师班组使用流程见 03_调控师操作手册.docx。",
        "附录C：飞书真实接入需企业自建应用、审批定义、多维表格权限和事件回调配置。",
    ])
    add_references(doc, [1, 2, 4, 5, 6, 9, 10, 11, 12, 13, 14])
    doc.add_heading("Market source and execution-price gate", level=2)
    table(
        doc,
        ["Source", "Use", "Execution authority"],
        [
            ["NBS production materials", "Public benchmark; ten-day publication", "No"],
            ["CZCE UR/SA", "Authorized futures trend reference", "No"],
            ["MOFCOM commodity price", "Periodic industry cross-check", "No"],
            ["Yuntu ERP / settlement", "Product, region, tax and freight aligned business price", "Yes after confirmation"],
        ],
        [2.0, 3.2, 1.2],
    )
    para(doc, "Public reference refresh only updates evidence cards. An unconfirmed or stale price cannot enter a new executable plan or economic ranking; the last valid enterprise version is retained and the business owner is notified.")
    return save(doc, "06_决赛完整方案文档.docx")


def build_feishu_module_doc():
    doc = setup_doc("飞书功能模块说明", "合成氨调度建议的确认、跟踪与班后记录")
    doc.add_heading("1. 模块定位", level=1)
    para(
        doc,
        "飞书承接调度建议送达、人工确认、任务跟踪和班后复盘。优化模块生成目标负荷、处置状态、触发字段、相对外售分配贡献和撤销条件；飞书保存负责人、会签状态、截止时间和执行结果。班后记录经专业人员审核后才进入案例库。",
    )
    doc.add_heading("1（一）围绕业务主线的飞书动作", level=2)
    table(
        doc,
        ["方向", "飞书动作", "沉淀结果"],
        [
            ["产供销协同", "卡片推送液氨去向、相对外售分配贡献、负荷调整、外采和停运评估；明确条件触发对应专业会签。", "调整原因、触发字段、确认人、测算值、实际结果和未采纳原因。"],
            ["跨装置协调", "审批后向合成、硝酸、罐区/调度、公辅四岗位分别派发动作单，全部接令后进入执行跟踪。", "动作单版本、接令人、时间、回复、异常和撤销条件。"],
            ["设备趋势异常", "趋势提示进入设备和班长任务，显示点位、持续时间和检查项，不直接开停机。", "趋势数据、通知对象、检查结果和处置记录。"],
            ["经验管理", "班后复盘写入飞书多维表格；Aily授权后只读取经审核知识源。", "带审核人、适用条件和版本的调度案例。"],
        ],
        [1.35, 3.1, 1.95],
    )
    doc.add_heading("2. 对三大命题痛点的回应", level=1)
    table(
        doc,
        ["痛点", "飞书承接动作", "解题价值"],
        [
            ["调度寻优难", "Aily授权后解释方案依据并标注来源；多维表格保存历史方案、未采纳原因和实际结果；事件回传进入待审核案例库。", "下一次同类场景可检索当时数据、判断和结果。"],
            ["指挥效率低", "互动卡片送达当班群，设备、公辅、库存等登记条件触发对应专业会签，任务清单记录执行与复盘。", "保存确认人、会签人、负责人、截止时间和关闭说明。"],
            ["知识传承难", "多维表格记录班次事实、班长备注、驳回原因、执行偏差和异常处置；智能问答只检索审核知识源。", "形成可追溯、可审核的班次案例和规则版本。"],
        ],
        [1.15, 2.75, 2.5],
    )
    doc.add_heading("3. 功能模块总览", level=1)
    table(
        doc,
        ["飞书能力", "合成氨业务对象", "首批字段/动作", "边界"],
        [
            ["互动卡片", "产供销、跨装置协调和设备趋势方案", "目标负荷、处置状态、触发字段、撤销条件、采纳/复核/驳回按钮", "只触发流程，不直接控制装置。"],
            ["审批", "设备趋势、公辅约束、液氨库存等明确条件会签", "会签人、装置、目标负荷、触发字段、现场影响、回滚条件", "命中演练停算线时不创建审批。"],
            ["多维表格Base", "班次事实与调度复盘", "订单、库存、设备健康、能耗窗口、采纳状态、执行偏差、未采纳原因", "作为知识事实层，保留字段来源和版本。"],
            ["飞书任务", "动作单接令与执行监督", "四岗位负责人、截止时间、接令状态、检查项、关闭说明", "接令不等于实际执行，不替代现场签字确认。"],
            ["飞书智能问答（Aily）", "调度员追问与新人学习", "结论、当前数值、边界、事实来源、责任人、复核时间", "待授权；无来源不回答，不绕过审批和安全规则。"],
            ["事件回传", "卡片点击、审批状态、任务关闭、复盘提交", "事件类型、防重复处理键、时间戳、操作者、处理结果", "待授权；必须验签、防重放并记录失败事件。"],
        ],
        [1.0, 1.45, 2.75, 1.2],
    )
    doc.add_heading("4. 当班闭环流程", level=1)
    numbers(
        doc,
        [
            "生成建议：调控师读取班次调度事实表，输出液氨分配、负荷调整和设备检查建议。",
            "卡片送达：将目标负荷、约束、相对外售分配贡献和撤销条件发送到合成氨当班调度群。",
            "条件确认：一般提示由调度员核对；设备、公辅、库存等登记条件成立时，由主管及对应专业会签。",
            "演练停算：触发后停止新目标、经济排序、动作单和审批，维持最近批准方案并转现场规程。",
            "动作单接令：审批通过后向合成主操、硝酸主操、罐区/调度和公辅调度分别派发任务，四岗位全部接令后才开放执行跟踪。",
            "执行监督：接令不等于实际执行；实际负荷、库存和能耗只从Historian或企业认可的只读事实源回传，无回传时保持为空。",
            "班后复盘：把采纳状态、未采纳原因、执行偏差和班长判断写入飞书多维表格。",
            "案例审核：生产、设备和安环专业审核复盘记录，再决定是否修订规则或纳入训练样本。",
        ],
    )
    doc.add_heading("4（一）条件驱动处置", level=2)
    table(doc, ["处置", "明确条件", "系统动作", "责任岗位"], CONDITION_ROUTING_ROWS, [1.0, 2.45, 2.15, 0.8])
    doc.add_heading("4（二）特殊停机取舍的飞书闭环", level=2)
    para(
        doc,
        "当流程确实无法维持连续时，调控师不直接下达停机指令，而是生成取舍卡片：比较空分、合成气轮机驱动系统与下游装置停运的安全影响、重启时间、废料损失和主流程影响。企业口述案例只作候选提示，卡片必须由班长、调度主管和相关专业会签，执行结果进入Base复盘库后仍需专家审核才能成为规则。",
    )
    add_dispatch_action_order(doc, "4（三）跨装置调度动作单", level=2)
    doc.add_heading("5. 数据表与接口契约", level=1)
    bullets(
        doc,
        [
            "字段样例：data/feishu_dispatch_review_template.csv 保存班次、场景、目标负荷、处置状态、触发字段、24小时相对外售分配贡献、采纳状态和复盘备注等字段。",
            "接口契约：data/feishu_integration_contract.json 保存卡片、审批、飞书多维表格、智能问答与事件回传的字段和防重复处理设计。",
            "趋势回放：data/compressor_trend_replay_sample.csv 验证观察、专业会签和演练停算流程，不作为企业设备阈值。",
            "字段矩阵：data/interface_field_matrix.csv 记录接口方向、频率、时间基准、质量规则、责任专业和超时动作。",
            "知识闭环：每条建议都必须保留 proposal_id、shift_id、scenario、evidence_completeness、approval_status、execution_delta 和 lesson_learned，便于后续回放和审计。",
        ],
    )
    doc.add_heading("6. 交付状态与授权说明", level=1)
    bullets(
        doc,
        [
            f"已验证：在线方案稿 {FEISHU_DOC_URL}",
            f"已验证：飞书多维表格样例库 {FEISHU_BASE_URL}",
            f"已验证：Task v2任务清单 {FEISHU_TASK_URL}",
            f"已验证：Task v2验收任务 {FEISHU_VERIFIED_TASK_URL}",
            "原型：互动卡片、负荷调整审批状态机、跨装置动作单门禁和事件回传字段契约。",
            "待授权：目标群机器人范围、审批定义approval_code、四岗位身份映射、Aily知识源、事件订阅、回调地址、密钥托管及企业系统接口。",
        ],
    )
    doc.add_heading("7. 安全边界", level=1)
    bullets(
        doc,
        [
            "平台不提供DCS/SIS写入路径；生产工况通过生产历史数据库或经批准的只读接口获取。",
            "数据过期、质量码无效、输入超出适用工况范围或限值接近时，停止给出可执行建议并转人工处理。",
            "所有回调事件使用签名校验、重放保护和幂等入库，避免重复写入或伪造操作。",
            "24小时指标称为相对外售分配贡献测算值，不称已实现收益；试点效果需剔除市场、检修和物流等外部因素。",
        ],
    )
    add_references(doc, [12, 13, 14, 15])
    return save(doc, "07_飞书功能模块说明.docx")


def build_demo_acceptance_doc():
    doc = setup_doc("决赛演示与验收附件", "五分钟演示、典型班次证据链、答辩口径与试点数据清单")
    doc.add_heading("1. 交付状态", level=1)
    table(doc, ["状态", "当前内容"], DELIVERY_STATUS_ROWS, [1.15, 5.25])
    para(doc, "演示中的价格、库存、负荷和设备趋势均为场景样例，不代表企业实时数据；现场数据、阈值、责任人和财务口径需在企业授权后确认。")

    doc.add_heading("2. 五分钟决赛演示脚本", level=1)
    table(doc, ["时间", "页面动作", "讲述要点", "当场证据"], DEMO_SCRIPT_ROWS, [0.8, 1.25, 2.85, 1.5])
    para(doc, "收束口径：平台不是替代操作员的自动控制器，而是把班次事实、调度判断、责任确认和执行结果串起来的工作方法。企业试点首先检验调度是否更快、更稳、更容易复盘。")
    doc.add_heading("2（一）条件驱动处置", level=2)
    table(doc, ["处置", "明确条件", "系统动作", "责任岗位"], CONDITION_ROUTING_ROWS, [1.0, 2.45, 2.15, 0.8])

    doc.add_heading("3. 典型班次证据链", level=1)
    para(doc, "演示班次设定为液氨库存趋紧、下游需求分化、硝酸行情走弱，同时压缩机趋势指标持续偏离。该设定用于展示跨系统判断和人工闭环，不作为现场工况结论。")
    table(doc, ["编号", "证据", "责任岗位", "保存位置", "状态"], EVIDENCE_CHAIN_ROWS, [0.55, 2.4, 1.25, 1.25, 0.95])
    para(doc, "E01至E04完整后才允许提交建议；明确条件成立时E05须完成对应专业会签；命中演练停算线时停止生成新目标、经济排序、动作单和审批，维持最近批准方案；E06四岗位未全部接令时不得进入执行跟踪；E07缺失时不得填写实际效果。")
    add_dispatch_action_order(doc, "3（一）跨装置调度动作单", level=2)

    doc.add_heading("4. 评委追问及事实回答", level=1)
    table(doc, ["追问", "事实回答"], JUDGE_QA_ROWS, [1.8, 4.6])

    doc.add_heading("5. 试点数据清单", level=1)
    table(doc, ["阶段", "数据域", "最小字段", "责任专业", "状态"], PILOT_DATA_ROWS, [0.65, 1.1, 2.8, 0.9, 0.95])
    table(doc, ["仓库样例", "用途", "关键字段", "使用边界"], DATA_SAMPLE_ROWS, [1.65, 1.15, 2.35, 1.25])
    para(doc, "所有字段同时记录来源系统、源时间、接收时间、单位、质量状态、版本和责任专业。密钥、个人联系方式和控制指令不进入分析数据集。")

    doc.add_heading("6. 试点验收口径", level=1)
    para(doc, "有效班次须同时具备冻结的原计划、合格输入、可追溯建议、完整人工确认、Historian实际回传和班后复盘。未执行方案可用于流程验收，但不进入效果统计。")
    table(doc, ["验收项", "口径", "判定"], ACCEPTANCE_ROWS, [1.05, 4.35, 1.0])
    table(doc, ["指标", "定义稿口径", "目标"], ACCEPTANCE_TARGET_ROWS, [1.35, 4.05, 1.0])
    para(doc, "四项目标是试点定义稿，不表示当前已达成。指标口径在30天数据校核后由企业冻结；没有同口径基线、真实执行回传或归因记录时，不给出节约金额、准确率或投资回报结论。")
    doc.add_heading("Market source demonstration", level=2)
    para(doc, "Show the official-reference cards and explain that the enterprise ERP or quotation version confirmed by the business owner is the only execution price. Public refresh, missing source, or stale data must not create a new executable plan.")
    return save(doc, "08_决赛演示与验收附件.docx")


def build_submission_summary(opening_filename="01_开题报告.docx"):
    doc = setup_doc("提交信息汇总", "在线表单填写索引、附件用途与体验链接")
    doc.add_heading("开题报告Part 1", level=1)
    para(doc, f"{PART1}（约{len(PART1)}字）")
    doc.add_heading("开题报告Part 2", level=1)
    para(doc, f"{PART2}（约{len(PART2)}字）")
    doc.add_heading("附件材料", level=1)
    table(
        doc,
        ["附件", "用途"],
        [
            [opening_filename, "开题报告定稿，集中回答命题理解、方案思路和资料依据。"],
            ["02_整体解决方案书.docx", "系统实施方案，说明业务流程、数据接口、约束模型、部署边界、试点步骤和验收方法。"],
            ["03_调控师操作手册.docx", "班组使用手册，说明调度员、班长和专业人员在班前、班中、班后的具体操作。"],
            ["04_参考文献与数据依据.docx", "证据来源与口径说明，记录事实来源、使用位置、适用边界及仍需企业确认的数据。"],
            ["05_方案创新与落地清单.docx", "交付状态与试点验收清单，区分已验证、原型、待授权和试点工作。"],
            ["06_决赛完整方案文档.docx", "参赛方案主文档，用典型班次串联问题、平台操作、飞书协同、结果复核和实施计划。"],
            ["07_飞书功能模块说明.docx", "飞书协同与接口说明，说明卡片、审批、Task v2、多维表格、智能问答和事件回传。"],
            ["08_决赛演示与验收附件.docx", "现场演示与答辩附件，包含五分钟脚本、班次证据链、跨装置动作单、事实问答、试点数据清单和验收口径。"],
        ],
        [2.1, 4.3],
    )
    doc.add_heading("链接材料", level=1)
    para(doc, "GitHub：https://github.com/zhijinDeng/synthesis-ammonia")
    para(doc, "本地平台入口：D:\\云图-合成氨-邓植斤\\index.html")
    para(doc, f"飞书在线完整方案稿：{FEISHU_DOC_URL}")
    para(doc, f"飞书Base调度复盘库原型：{FEISHU_BASE_URL}")
    para(doc, f"飞书任务清单：{FEISHU_TASK_URL}")
    para(doc, f"飞书Task v2验收任务：{FEISHU_VERIFIED_TASK_URL}")
    return save(doc, "00_提交信息汇总.docx")


def main():
    opening_path = build_opening_report()
    paths = [
        build_submission_summary(opening_path.name),
        opening_path,
        build_solution_doc(),
        build_manual_doc(),
        build_reference_doc(),
        build_innovation_doc(),
        build_final_doc(),
        build_feishu_module_doc(),
        build_demo_acceptance_doc(),
    ]
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
